#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
generate_report.py

Builds "Design and Implementation of a 4x4 Weight-Stationary Systolic Array
using Multiply-Accumulate (MAC) Units" as a Microsoft Word (.docx) report.

Run:  python generate_report.py
Out:  Systolic_Array_Project_Report.docx

After opening the document in Word, press Ctrl+A then F9 (choose "Update entire
table") to populate the Table of Contents and the List of Figures/Tables.
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT, WD_BREAK
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUTPUT = "Systolic_Array_Project_Report.docx"

BODY_FONT = "Times New Roman"
MONO_FONT = "Consolas"
ACCENT = RGBColor(0x1F, 0x3B, 0x73)      # deep blue for headings
GREY = RGBColor(0x44, 0x44, 0x44)

# ---------------------------------------------------------------------------
# Low-level XML helpers
# ---------------------------------------------------------------------------


def _shade(element, fill):
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill)
    element.append(shd)


def shade_paragraph(par, fill="F2F4F8"):
    _shade(par._p.get_or_add_pPr(), fill)


def shade_cell(cell, fill="1F3B73"):
    _shade(cell._tc.get_or_add_tcPr(), fill)


def box_paragraph(par, colour="9AA5B1", size=6):
    """Draw a thin box around a paragraph (used for figure placeholders)."""
    pbdr = OxmlElement("w:pBdr")
    for edge in ("top", "left", "bottom", "right"):
        e = OxmlElement(f"w:{edge}")
        e.set(qn("w:val"), "single")
        e.set(qn("w:sz"), str(size))
        e.set(qn("w:space"), "6")
        e.set(qn("w:color"), colour)
        pbdr.append(e)
    par._p.get_or_add_pPr().append(pbdr)


def add_field(par, instruction):
    """Insert a Word field code (used for TOC, page numbers, captions)."""
    r1 = par.add_run()
    fld = OxmlElement("w:fldChar")
    fld.set(qn("w:fldCharType"), "begin")
    r1._r.append(fld)

    r2 = par.add_run()
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = instruction
    r2._r.append(instr)

    r3 = par.add_run()
    fld = OxmlElement("w:fldChar")
    fld.set(qn("w:fldCharType"), "separate")
    r3._r.append(fld)

    r4 = par.add_run("(update field: Ctrl+A, F9)")
    r4.font.size = Pt(9)
    r4.font.color.rgb = GREY

    r5 = par.add_run()
    fld = OxmlElement("w:fldChar")
    fld.set(qn("w:fldCharType"), "end")
    r5._r.append(fld)


def add_page_number_footer(section):
    par = section.footer.paragraphs[0]
    par.alignment = WD_ALIGN_PARAGRAPH.CENTER
    # simple PAGE field
    r1 = par.add_run()
    f = OxmlElement("w:fldChar"); f.set(qn("w:fldCharType"), "begin"); r1._r.append(f)
    r2 = par.add_run()
    i = OxmlElement("w:instrText"); i.set(qn("xml:space"), "preserve"); i.text = " PAGE "
    r2._r.append(i)
    r3 = par.add_run()
    f = OxmlElement("w:fldChar"); f.set(qn("w:fldCharType"), "end"); r3._r.append(f)
    for r in par.runs:
        r.font.name = BODY_FONT
        r.font.size = Pt(10)


# ---------------------------------------------------------------------------
# Document / style setup
# ---------------------------------------------------------------------------

doc = Document()

sec = doc.sections[0]
sec.page_width = Inches(8.27)      # A4
sec.page_height = Inches(11.69)
sec.left_margin = Inches(1.25)
sec.right_margin = Inches(1.0)
sec.top_margin = Inches(1.0)
sec.bottom_margin = Inches(1.0)

st = doc.styles["Normal"]
st.font.name = BODY_FONT
st.font.size = Pt(12)
st.element.rPr.rFonts.set(qn("w:eastAsia"), BODY_FONT)
st.paragraph_format.line_spacing = 1.5
st.paragraph_format.space_after = Pt(6)
st.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

for name, size, colour in (("Heading 1", 17, ACCENT),
                           ("Heading 2", 14, ACCENT),
                           ("Heading 3", 12.5, ACCENT)):
    s = doc.styles[name]
    s.font.name = BODY_FONT
    s.font.size = Pt(size)
    s.font.bold = True
    s.font.color.rgb = colour
    s.paragraph_format.space_before = Pt(14)
    s.paragraph_format.space_after = Pt(8)
    s.paragraph_format.line_spacing = 1.2
    s.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    s.paragraph_format.keep_with_next = True

cap = doc.styles["Caption"]
cap.font.name = BODY_FONT
cap.font.size = Pt(10)
cap.font.bold = False
cap.font.italic = True
cap.font.color.rgb = GREY

# Separate caption styles for figures and tables so that Word can build a
# List of Figures and a List of Tables from them (TOC \t "<style>,1").
for _cap_style in ("Figure Caption", "Table Caption"):
    s = doc.styles.add_style(_cap_style, WD_STYLE_TYPE.PARAGRAPH)
    s.base_style = doc.styles["Caption"]
    s.font.name = BODY_FONT
    s.font.size = Pt(10)
    s.font.italic = True
    s.font.color.rgb = GREY
    s.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    s.paragraph_format.line_spacing = 1.15
    s.quick_style = True

# ---------------------------------------------------------------------------
# Content helpers
# ---------------------------------------------------------------------------

_fig_no = {}
_tab_no = {}
_eq_no = {}
CHAPTER = [0]


def h1(text):
    CHAPTER[0] += 1
    doc.add_page_break()
    p = doc.add_heading(text, level=1)
    return p


def h1_plain(text):
    """Front/back matter heading that does not bump the chapter counter."""
    doc.add_page_break()
    return doc.add_heading(text, level=1)


def h2(text):
    return doc.add_heading(text, level=2)


def h3(text):
    return doc.add_heading(text, level=3)


def p(text, style=None, italic=False, align=None):
    par = doc.add_paragraph(style=style)
    run = par.add_run(text)
    run.italic = italic
    if align is not None:
        par.alignment = align
    return par


def bullets(items, numbered=False):
    style = "List Number" if numbered else "List Bullet"
    for it in items:
        par = doc.add_paragraph(style=style)
        par.paragraph_format.line_spacing = 1.3
        par.paragraph_format.space_after = Pt(3)
        if isinstance(it, tuple):
            r = par.add_run(it[0])
            r.bold = True
            par.add_run(" — " + it[1])
        else:
            par.add_run(it)


def equation(text, label=None):
    """Centred display equation with a right-aligned equation number."""
    ch = CHAPTER[0]
    _eq_no[ch] = _eq_no.get(ch, 0) + 1
    num = f"({ch}.{_eq_no[ch]})"
    par = doc.add_paragraph()
    par.paragraph_format.line_spacing = 1.2
    par.paragraph_format.space_before = Pt(8)
    par.paragraph_format.space_after = Pt(8)
    par.paragraph_format.tab_stops.add_tab_stop(Inches(2.6), WD_TAB_ALIGNMENT.CENTER)
    par.paragraph_format.tab_stops.add_tab_stop(Inches(5.6), WD_TAB_ALIGNMENT.RIGHT)
    run = par.add_run("\t" + text)
    run.italic = True
    run.font.size = Pt(12.5)
    par.add_run("\t" + num)
    return num


def figure(art, caption, note=None):
    """A boxed figure placeholder containing a schematic, plus its caption."""
    ch = CHAPTER[0]
    _fig_no[ch] = _fig_no.get(ch, 0) + 1
    num = f"{ch}.{_fig_no[ch]}"

    par = doc.add_paragraph()
    par.alignment = WD_ALIGN_PARAGRAPH.CENTER
    par.paragraph_format.line_spacing = 1.0
    par.paragraph_format.space_before = Pt(10)
    par.paragraph_format.space_after = Pt(2)
    run = par.add_run(art)
    run.font.name = MONO_FONT
    run.font.size = Pt(8.5)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), MONO_FONT)
    box_paragraph(par)
    shade_paragraph(par, "FAFBFC")

    cpar = doc.add_paragraph(style="Figure Caption")
    cpar.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cpar.paragraph_format.space_after = Pt(12)
    cr = cpar.add_run(f"Figure {num}: {caption}")
    cr.italic = True
    if note:
        npar = doc.add_paragraph(style="Caption")
        npar.alignment = WD_ALIGN_PARAGRAPH.CENTER
        npar.add_run(note).italic = True
    return num


def table(caption, headers, rows, widths=None):
    ch = CHAPTER[0]
    _tab_no[ch] = _tab_no.get(ch, 0) + 1
    num = f"{ch}.{_tab_no[ch]}"

    cpar = doc.add_paragraph(style="Table Caption")
    cpar.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cpar.paragraph_format.space_before = Pt(10)
    cpar.add_run(f"Table {num}: {caption}").italic = True

    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER

    hdr = t.rows[0].cells
    for i, htxt in enumerate(headers):
        shade_cell(hdr[i])
        par = hdr[i].paragraphs[0]
        par.alignment = WD_ALIGN_PARAGRAPH.CENTER
        par.paragraph_format.line_spacing = 1.0
        par.paragraph_format.space_after = Pt(2)
        r = par.add_run(htxt)
        r.bold = True
        r.font.size = Pt(10.5)
        r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            par = cells[i].paragraphs[0]
            par.alignment = (WD_ALIGN_PARAGRAPH.LEFT if i == 0
                             else WD_ALIGN_PARAGRAPH.CENTER)
            par.paragraph_format.line_spacing = 1.0
            par.paragraph_format.space_after = Pt(2)
            r = par.add_run(str(val))
            r.font.size = Pt(10.5)

    if widths:
        for row in t.rows:
            for i, w in enumerate(widths):
                row.cells[i].width = Inches(w)

    doc.add_paragraph().paragraph_format.space_after = Pt(6)
    return num


def summary(text):
    h2("Chapter Summary")
    par = doc.add_paragraph()
    par.paragraph_format.left_indent = Inches(0.15)
    par.paragraph_format.right_indent = Inches(0.15)
    par.paragraph_format.space_before = Pt(4)
    shade_paragraph(par, "F2F4F8")
    run = par.add_run(text)
    run.font.size = Pt(11.5)


# ===========================================================================
# TITLE PAGE
# ===========================================================================

for _ in range(2):
    doc.add_paragraph()

t = doc.add_paragraph()
t.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = t.add_run("DESIGN AND IMPLEMENTATION OF A 4×4\nWEIGHT-STATIONARY SYSTOLIC ARRAY\nUSING MULTIPLY-ACCUMULATE (MAC) UNITS")
r.bold = True
r.font.size = Pt(22)
r.font.color.rgb = ACCENT
t.paragraph_format.line_spacing = 1.3

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = sub.add_run("A Hardware Accelerator for Matrix Multiplication in\nDeep Neural Network Inference")
r.italic = True
r.font.size = Pt(13.5)
sub.paragraph_format.space_before = Pt(18)

doc.add_paragraph()
line = doc.add_paragraph()
line.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = line.add_run("─" * 46)
r.font.color.rgb = ACCENT

for _ in range(2):
    doc.add_paragraph()

for txt, size, bold in (("PROJECT REPORT", 13, True),
                        ("Submitted in partial fulfilment of the requirements\n"
                         "for the award of the degree of", 11.5, False),
                        ("BACHELOR OF ENGINEERING", 13, True),
                        ("in", 11.5, False),
                        ("ELECTRONICS AND COMMUNICATION ENGINEERING", 12.5, True)):
    par = doc.add_paragraph()
    par.alignment = WD_ALIGN_PARAGRAPH.CENTER
    par.paragraph_format.space_after = Pt(8)
    run = par.add_run(txt)
    run.bold = bold
    run.font.size = Pt(size)

for _ in range(3):
    doc.add_paragraph()

for txt in ("Submitted by", "[ Student Name ]    [ Register Number ]",
            "Under the guidance of", "[ Guide Name, Designation ]"):
    par = doc.add_paragraph()
    par.alignment = WD_ALIGN_PARAGRAPH.CENTER
    par.paragraph_format.space_after = Pt(4)
    run = par.add_run(txt)
    run.font.size = Pt(11.5)
    if txt.startswith("["):
        run.bold = True

for _ in range(3):
    doc.add_paragraph()

for txt, bold in (("[ DEPARTMENT OF ELECTRONICS AND COMMUNICATION ENGINEERING ]", True),
                  ("[ NAME OF THE INSTITUTION ]", True),
                  ("[ MONTH, YEAR ]", False)):
    par = doc.add_paragraph()
    par.alignment = WD_ALIGN_PARAGRAPH.CENTER
    par.paragraph_format.space_after = Pt(4)
    run = par.add_run(txt)
    run.bold = bold
    run.font.size = Pt(12)

add_page_number_footer(sec)

# ===========================================================================
# ABSTRACT
# ===========================================================================

h1_plain("Abstract")

p("Matrix multiplication is the computational heart of modern artificial "
  "intelligence. A single inference pass through a deep neural network may "
  "require billions of multiply-accumulate operations, and the overwhelming "
  "majority of them belong to dense matrix-matrix or matrix-vector products. "
  "General-purpose processors execute these products correctly but "
  "inefficiently: for every useful arithmetic operation they must also fetch "
  "an instruction, decode it, check dependencies, read a register file and "
  "write a result back. The arithmetic is cheap; the movement of data and the "
  "administration of control are not.")

p("This project addresses that inefficiency by designing and implementing a "
  "4×4 weight-stationary systolic array built from sixteen Multiply-Accumulate "
  "(MAC) Processing Elements. In a systolic array, data is pumped rhythmically "
  "through a regular mesh of simple processors, each of which performs one "
  "multiply-accumulate per clock cycle and then hands its results to its "
  "neighbours. Because every operand entering the array is reused by many "
  "processing elements before it is discarded, the array performs a large "
  "amount of arithmetic for a very small amount of memory traffic.")

p("The weight-stationary dataflow is adopted for this design. Each processing "
  "element stores one weight for the entire duration of a computation, "
  "activations stream horizontally across the array, and partial sums "
  "accumulate vertically down each column. The report develops this "
  "architecture from first principles: it begins with the definition of "
  "multiplication, builds up to matrix multiplication and its cost, "
  "establishes why conventional processors are inadequate for it, introduces "
  "systolic arrays as the architectural answer, describes the internal "
  "structure of a single processing element in complete detail, justifies the "
  "weight-stationary choice against alternative dataflows, assembles the "
  "complete 4×4 array with a full clock-by-clock timing analysis, and finally "
  "shows how the same principles are applied at industrial scale in Google's "
  "Tensor Processing Unit.")

p("Analysis of the resulting design shows that the 4×4 array completes a "
  "4×4 × 4×4 matrix product in 11 clock cycles after a 4-cycle weight-load "
  "phase, against 64 cycles for a scalar processor issuing one MAC per cycle. "
  "In sustained operation, with a single set of weights reused across many "
  "activation rows, the array delivers 16 MAC operations per clock cycle — "
  "one per processing element — while requiring only one weight fetch per "
  "processing element for the entire workload.")

par = doc.add_paragraph()
par.paragraph_format.space_before = Pt(14)
r = par.add_run("Keywords: ")
r.bold = True
par.add_run("Systolic array, weight-stationary dataflow, multiply-accumulate "
            "(MAC), processing element, matrix multiplication, hardware "
            "accelerator, deep neural network inference, Tensor Processing "
            "Unit (TPU), data reuse, pipelining.")

# ===========================================================================
# TABLE OF CONTENTS
# ===========================================================================

h1_plain("Table of Contents")
par = doc.add_paragraph()
add_field(par, r'TOC \o "1-3" \h \z \u')

h1_plain("List of Figures")
par = doc.add_paragraph()
add_field(par, r'TOC \h \z \t "Figure Caption,1"')

doc.add_paragraph()
doc.add_heading("List of Tables", level=1)
par = doc.add_paragraph()
add_field(par, r'TOC \h \z \t "Table Caption,1"')

# ===========================================================================
# CHAPTER 1 — INTRODUCTION
# ===========================================================================

h1("Chapter 1: Introduction")

p("Every computing machine ever built exists to answer a question that a human "
  "being could not answer quickly enough by hand. The history of computation is "
  "therefore the history of arithmetic being pushed into machinery: first "
  "counting, then addition, then multiplication, and eventually the enormous "
  "structured multiplications that underlie modern artificial intelligence. "
  "This chapter follows that progression, because the architecture developed "
  "in the remainder of this report is a direct consequence of it.")

h2("1.1 Introduction to Multiplication")

p("Multiplication is, at its root, repeated addition. The product 4 × 3 is "
  "nothing more than the sum 4 + 4 + 4. This definition is intuitive but "
  "computationally poor: multiplying two 8-bit numbers by repeated addition "
  "could require up to 255 additions. Digital hardware therefore implements "
  "multiplication using the shift-and-add method, which exploits the "
  "positional structure of binary numbers.")

p("For two n-bit unsigned binary numbers A and B, the product is expressed as "
  "the weighted sum of shifted copies of the multiplicand:")

equation("P = A × B = Σ (from i = 0 to n−1) of  bᵢ · A · 2ⁱ")

p("where bᵢ is the i-th bit of B. Each term is either zero (when bᵢ = 0) "
  "or a left-shifted copy of A (when bᵢ = 1). A hardware multiplier "
  "generates these n partial products in parallel and reduces them to a single "
  "result using an adder tree. This is why a multiplier occupies far more "
  "silicon area and consumes far more energy than an adder: it is, in effect, "
  "an entire array of adders.")

p("Two practical consequences follow, and both shape the design in this report. "
  "First, the multiplier is the most expensive arithmetic resource on the chip, "
  "so an architecture is judged largely by how busy it keeps its multipliers. "
  "Second, the product of two n-bit numbers requires 2n bits to represent "
  "exactly. When many such products are summed, further bits are needed to "
  "prevent overflow. This is the reason the accumulator in the processing "
  "element designed later in this report is 32 bits wide while its operands are "
  "only 8 bits wide.")

h3("1.1.1 Signed Multiplication")

p("Neural network weights and activations are signed quantities. In two's "
  "complement representation, an n-bit signed number has the value")

equation("A = −aₙ₋₁ · 2ⁿ⁻¹ + Σ (from i = 0 to n−2) of  aᵢ · 2ⁱ")

p("The most significant bit carries negative weight. A signed multiplier must "
  "therefore sign-extend its operands before forming partial products; treating "
  "a negative operand as unsigned produces a result that is wrong by a large "
  "power of two. For 8-bit signed operands the value range is −128 to +127, and "
  "the extreme product (−128) × (−128) = +16 384 requires 16 bits — a detail "
  "that must be handled correctly in hardware and is verified explicitly in "
  "this project.")

h2("1.2 Introduction to Matrix Multiplication")

p("A matrix is a rectangular array of numbers. Matrix multiplication is the "
  "operation that combines two matrices into a third, and it is defined so that "
  "it represents the composition of two linear transformations. If A is an "
  "m × k matrix and B is a k × n matrix, their product C = A × B is an m × n "
  "matrix given by")

equation("C = A × B")

p("with each element of C defined as the dot product of a row of A with a "
  "column of B:")

equation("Cᵢⱼ = Σ (from k = 0 to K−1) of  Aᵢₖ · Bₖⱼ")

p("Two properties of this definition deserve emphasis because the entire "
  "architecture of a systolic array is built on them. First, the inner "
  "computation is a sum of products — a multiply-accumulate. Second, every "
  "element of A is used in n different dot products (once for each column of "
  "B), and every element of B is used in m different dot products. The data has "
  "enormous inherent reuse. A good architecture exploits that reuse; a poor "
  "architecture re-fetches the same value from memory over and over.")

h3("1.2.1 A Complete 2×2 Worked Example")

p("Consider the two matrices")

figure(
    "        ⎡ 1   2 ⎤            ⎡ 5   6 ⎤\n"
    "  A  =  ⎢       ⎥      B  =  ⎢       ⎥\n"
    "        ⎣ 3   4 ⎦            ⎣ 7   8 ⎦",
    "Two 2×2 matrices used for the worked example.")

p("Each element of the product is computed as a row-times-column dot product:")

equation("C₀₀ = A₀₀·B₀₀ + A₀₁·B₁₀ = (1)(5) + (2)(7) = 5 + 14 = 19")
equation("C₀₁ = A₀₀·B₀₁ + A₀₁·B₁₁ = (1)(6) + (2)(8) = 6 + 16 = 22")
equation("C₁₀ = A₁₀·B₀₀ + A₁₁·B₁₀ = (3)(5) + (4)(7) = 15 + 28 = 43")
equation("C₁₁ = A₁₀·B₀₁ + A₁₁·B₁₁ = (3)(6) + (4)(8) = 18 + 32 = 50")

p("giving the result")

figure(
    "        ⎡ 19   22 ⎤\n"
    "  C  =  ⎢         ⎥\n"
    "        ⎣ 43   50 ⎦",
    "Result of the 2×2 matrix multiplication C = A × B.")

p("Observe the structure of the work: four output elements, each requiring two "
  "multiplications and one addition — eight multiplications and four additions "
  "in total. Observe also the reuse: the value A₀₀ = 1 was needed for "
  "both C₀₀ and C₀₁, and the value B₀₀ = 5 was "
  "needed for both C₀₀ and C₁₀. Every operand was used "
  "exactly twice. In a 4×4 product every operand is used four times; in the "
  "1024×1024 products typical of neural networks, every operand is used 1024 "
  "times. Capturing that reuse in hardware, rather than paying for it in memory "
  "traffic, is the single most important idea in this report.")

h2("1.3 Why Matrix Multiplication is Important")

p("Matrix multiplication is not merely one operation among many; it is the "
  "operation into which an extraordinary range of computational problems is "
  "translated. Once a problem has been expressed as a matrix product, it "
  "inherits every optimisation ever devised for matrix products — which is "
  "precisely why so much effort is invested in making that one operation fast.")

h3("1.3.1 Artificial Intelligence and Machine Learning")

p("A fully connected neural network layer computes")

equation("y = f(W · x + b)")

p("where x is the input activation vector, W the weight matrix, b the bias "
  "vector and f a non-linear activation function such as ReLU. The dominant "
  "cost is the matrix-vector product W · x. When a batch of inputs is processed "
  "together, the batch is stacked into a matrix X and the layer computes "
  "W · X — a full matrix-matrix product. Convolutional layers, although "
  "conceptually different, are almost universally implemented as matrix "
  "multiplications after an im2col transformation, and the attention mechanism "
  "at the core of transformer models consists of three large matrix products "
  "per layer. Across virtually all modern network architectures, between 80 % "
  "and 95 % of the arithmetic operations in an inference pass are "
  "multiply-accumulate operations belonging to a matrix product.")

h3("1.3.2 Digital Signal Processing")

p("Filtering, correlation and transforms are matrix operations. A finite "
  "impulse response filter is a dot product between a coefficient vector and a "
  "sliding window of samples; the Discrete Fourier Transform is a multiplication "
  "by a fixed n × n matrix of twiddle factors; beam-forming in radar and "
  "wireless communication multiplies an antenna sample vector by a steering "
  "matrix. All of these reduce to sums of products evaluated at very high rates.")

h3("1.3.3 Computer Graphics")

p("Every vertex in a three-dimensional scene is transformed by a 4×4 "
  "homogeneous matrix that encodes rotation, scaling, translation and "
  "perspective projection. A frame containing a million vertices therefore "
  "requires a million 4×4 matrix-vector products. The prominence of the 4×4 "
  "matrix in graphics is one historical reason why 4×4 is a natural "
  "demonstration size for a systolic array.")

h3("1.3.4 Robotics and Control")

p("Robot kinematics chains together homogeneous transformation matrices to "
  "relate joint angles to end-effector position. State estimation algorithms "
  "such as the Kalman filter perform matrix multiplications and inversions at "
  "every time step. In both cases the computation is on a hard real-time "
  "deadline, so the throughput of the matrix engine directly determines the "
  "achievable control-loop frequency.")

h3("1.3.5 Scientific and Engineering Computation")

p("Finite element analysis, computational fluid dynamics, structural "
  "simulation, quantum chemistry and weather modelling all reduce ultimately to "
  "the solution of large linear systems, and the inner kernel of nearly every "
  "such solver is a dense matrix multiplication. The LINPACK benchmark used to "
  "rank the world's supercomputers is, in essence, a measurement of matrix "
  "multiplication throughput.")

table("Application domains dominated by matrix multiplication.",
      ["Domain", "Typical Operation", "Why It Matters"],
      [["Deep learning inference", "W · x per layer", "80–95 % of all operations"],
       ["Deep learning training", "Forward + gradient products", "Weeks of compute time"],
       ["Digital signal processing", "FIR, FFT, beam-forming", "Real-time sample rates"],
       ["Computer graphics", "4×4 vertex transforms", "Millions of vertices per frame"],
       ["Robotics and control", "Kinematics, Kalman filter", "Hard real-time deadlines"],
       ["Scientific computing", "Dense linear algebra", "Defines supercomputer ranking"]],
      widths=[2.1, 1.9, 2.0])

summary(
    "Multiplication in hardware is implemented as a shift-and-add reduction of "
    "partial products, making the multiplier the most expensive arithmetic unit "
    "on a chip and requiring accumulators wider than the operands. Matrix "
    "multiplication, defined by Cij = Σ Aik·Bkj, builds directly on the "
    "multiply-accumulate primitive and exhibits massive inherent data reuse: in "
    "an n×n product every operand participates in n different dot products. "
    "Because matrix multiplication dominates artificial intelligence, signal "
    "processing, graphics, robotics and scientific computing, accelerating this "
    "single operation accelerates an enormous fraction of all computation. The "
    "next chapter introduces the architectural structure — the systolic array — "
    "that is specifically designed to exploit that reuse.")

# ===========================================================================
# CHAPTER 2 — SYSTOLIC ARRAY FUNDAMENTALS
# ===========================================================================

h1("Chapter 2: Systolic Array Fundamentals")

p("Chapter 1 established that matrix multiplication consists of a very large "
  "number of multiply-accumulate operations performed on data with a very "
  "regular reuse pattern. This chapter introduces the architecture designed "
  "precisely for such a workload.")

h2("2.1 What is a Systolic Array?")

p("A systolic array is a network of identical, simple processing elements "
  "arranged in a regular geometric pattern — usually a one- or two-dimensional "
  "mesh — in which data flows rhythmically from processing element to "
  "neighbouring processing element under the control of a single global clock. "
  "Each processing element performs a small fixed computation on the data that "
  "arrives at its inputs, stores a result if required, and passes data onward "
  "to its neighbours. Crucially, a processing element communicates only with "
  "its immediate neighbours; there is no global bus, no shared register file "
  "and no memory hierarchy inside the array.")

p("The concept was introduced by H. T. Kung and Charles E. Leiserson at "
  "Carnegie Mellon University in 1978, motivated by the economics of "
  "Very Large Scale Integration. Their observation was that in VLSI, "
  "computation is cheap and communication is expensive: transistors shrink and "
  "multiply rapidly, but long wires do not become proportionately faster, and "
  "the energy required to move a bit across a chip does not fall as quickly as "
  "the energy required to compute with it. An architecture that replaces long "
  "global wires with short local wires between neighbours therefore scales far "
  "better than one built around a central bus.")

p("The defining characteristics of a systolic array are:")

bullets([
    ("Regularity", "all processing elements are identical, so the layout is "
     "generated by replicating a single cell; verification effort and design "
     "risk are correspondingly small."),
    ("Locality", "each processing element communicates only with its adjacent "
     "neighbours, so all interconnect is short, fast and low-energy."),
    ("Rhythmic operation", "data advances by exactly one processing element per "
     "clock cycle, making the timing entirely deterministic."),
    ("Massive parallelism", "every processing element computes on every cycle, "
     "so an n × n array performs n² operations per cycle."),
    ("High data reuse", "an operand entering the array at its boundary is used "
     "by many processing elements before it leaves, so external memory is "
     "accessed once for many computations."),
])

h2("2.2 Why is it Called \"Systolic\"?")

p("The name is a deliberate biological metaphor. In cardiology, systole is the "
  "contraction phase of the heartbeat, during which the heart muscle pumps "
  "blood out through the arteries into the body. The pulse is regular, "
  "rhythmic and drives blood through a fixed network of vessels; each organ "
  "extracts what it needs as the blood passes through, and the blood continues "
  "on its way.")

p("Kung and Leiserson saw exactly this structure in their architecture. The "
  "global clock is the heartbeat. Data is the blood. The processing elements "
  "are the organs, each taking data as it passes, doing useful work with it, "
  "and passing it along. Data is \"pumped\" through the array in a regular "
  "rhythm rather than being fetched on demand. Just as blood circulates "
  "through the body many times before being replenished, an operand pumped "
  "into a systolic array is used many times before being discarded. The "
  "metaphor is not decorative: it captures the essential contrast with a "
  "conventional processor, where data is pulled from memory on demand and "
  "returned immediately after a single use.")

figure(
    "        MEMORY                                MEMORY\n"
    "          │                                     │\n"
    "          ▼                                     ▼\n"
    "        ┌───┐   fetch                        ┌─────┐\n"
    "        │CPU│◄────────►                      │ PE  │──►┌─────┐──►┌─────┐\n"
    "        └───┘   store                        └─────┘   │ PE  │   │ PE  │\n"
    "          ▲                                     │      └─────┘   └─────┘\n"
    "          │                                     ▼         │          │\n"
    "        MEMORY                              ┌─────┐       ▼          ▼\n"
    "                                            │ PE  │──►┌─────┐──►┌─────┐\n"
    "                                            └─────┘   │ PE  │   │ PE  │\n"
    "                                               │      └─────┘   └─────┘\n"
    "                                               ▼         │          │\n"
    "   (a) von Neumann: one operand,               ▼         ▼          ▼\n"
    "       one memory access, one result       (b) Systolic: one memory access,\n"
    "                                               many computations",
    "Data movement in a conventional processor compared with a systolic array. "
    "In (a) every operand makes a round trip to memory; in (b) an operand "
    "entering at the boundary is consumed by many processing elements.")

h2("2.3 Working Principle")

p("The operating principle of a systolic array can be stated in one sentence: "
  "replace repeated memory accesses with repeated use of data already in "
  "motion. This is achieved by three mechanisms working together.")

h3("2.3.1 Pipelining in Two Dimensions")

p("A conventional pipeline is one-dimensional: an instruction advances through "
  "fetch, decode, execute and write-back stages. A systolic array is a "
  "two-dimensional pipeline. Activations advance horizontally while partial "
  "sums advance vertically, and every processing element is simultaneously a "
  "stage of one horizontal pipeline and one vertical pipeline. After an initial "
  "fill period, every stage of every pipeline is busy on every cycle.")

h3("2.3.2 Local Communication Only")

p("Because a processing element reads only from its immediate neighbours, the "
  "longest wire in the array is the distance between two adjacent cells. This "
  "keeps the critical path short and independent of array size, so the clock "
  "frequency does not degrade as the array grows. It also means that the "
  "array can be laid out physically as a simple tile-and-abut replication of "
  "one cell.")

h3("2.3.3 Distributed Storage")

p("Intermediate results never return to memory. A partial sum is produced by "
  "one processing element and consumed by its neighbour on the very next "
  "cycle; it exists only in a pipeline register. Only the final result of a "
  "complete dot product is written out of the array. In a 4×4 array, three of "
  "every four intermediate values never leave the array at all.")

h2("2.4 Data Flow")

p("In the two-dimensional array used for matrix multiplication, three distinct "
  "data streams move through the mesh, each in its own direction:")

bullets([
    ("Weights (stationary)", "loaded once into each processing element before "
     "computation begins and held there for the entire operation."),
    ("Activations (horizontal)", "enter at the left edge and propagate "
     "rightward, advancing one processing element per clock cycle."),
    ("Partial sums (vertical)", "enter at the top edge, accumulate as they "
     "descend, and emerge completed at the bottom edge."),
])

p("Because the streams travel in different directions, an operand meets a "
  "different partner in every processing element it visits. This is the "
  "geometric mechanism by which a single activation contributes to several "
  "different output elements without ever being fetched twice.")

figure(
    "                   partial sums enter (psum_in)\n"
    "                    │        │        │        │\n"
    "                    ▼        ▼        ▼        ▼\n"
    "   activations   ┌──────┐ ┌──────┐ ┌──────┐ ┌──────┐\n"
    "   enter ───────►│ PE   ├►│ PE   ├►│ PE   ├►│ PE   ├──► discarded\n"
    "                 └──┬───┘ └──┬───┘ └──┬───┘ └──┬───┘\n"
    "                    ▼        ▼        ▼        ▼\n"
    "                 ┌──────┐ ┌──────┐ ┌──────┐ ┌──────┐\n"
    "        ───────► │ PE   ├►│ PE   ├►│ PE   ├►│ PE   ├──►\n"
    "                 └──┬───┘ └──┬───┘ └──┬───┘ └──┬───┘\n"
    "                    ▼        ▼        ▼        ▼\n"
    "                completed results leave (psum_out)",
    "The three data streams of a weight-stationary systolic array: stationary "
    "weights held inside each processing element, activations flowing left to "
    "right, and partial sums flowing top to bottom.")

h2("2.5 The MAC Equation")

p("The single arithmetic operation performed by every processing element on "
  "every clock cycle is the multiply-accumulate:")

equation("Psum = Psum + (A × B)")

p("In the notation used throughout the remainder of this report, the "
  "processing element at array position (i, j) computes")

equation("psum_out = psum_in + (weight × a_in)")

p("This is the complete functional specification of a processing element. "
  "Everything else — the registers, the control signals, the interconnect — "
  "exists only to deliver the right operands to this equation at the right "
  "time and to move the result to the right place.")

p("The power of the architecture lies in replication. One processing element "
  "performs one MAC per cycle. A 4×4 array performs sixteen MACs per cycle. "
  "The 256×256 array inside Google's first Tensor Processing Unit performs "
  "65 536 MACs per cycle. The arithmetic in each cell is trivial; the "
  "architecture is what turns triviality into throughput.")

table("Terminology used consistently throughout this report.",
      ["Term", "Symbol / Signal", "Meaning"],
      [["Activation", "a_in, a_out", "Input data element, flows horizontally"],
       ["Weight", "weight_q", "Stationary coefficient held inside a PE"],
       ["Partial sum", "psum_in, psum_out", "Running accumulation, flows vertically"],
       ["MAC", "—", "One multiply followed by one accumulate"],
       ["PE", "—", "Processing Element: one MAC unit plus registers"],
       ["Latency", "—", "Cycles from first input to first output"],
       ["Throughput", "—", "Results produced per cycle in steady state"]],
      widths=[1.6, 1.7, 2.7])

summary(
    "A systolic array is a regular mesh of identical processing elements that "
    "communicate only with their immediate neighbours and through which data is "
    "pumped rhythmically under a global clock — the property that gives the "
    "architecture its name, borrowed from the systolic contraction of the "
    "heart. Its advantages follow from replacing expensive global communication "
    "with cheap local communication and from reusing every operand many times "
    "after a single memory fetch. Each processing element performs the "
    "elementary operation Psum = Psum + (A × B) once per clock cycle, and "
    "throughput is obtained purely by replicating that element. The next "
    "chapter quantifies exactly how much work a matrix multiplication requires "
    "and how much of it the systolic organisation eliminates.")

# ===========================================================================
# CHAPTER 3 — MATRIX MULTIPLICATION
# ===========================================================================

h1("Chapter 3: Matrix Multiplication")

p("Having introduced both the operation and the architecture, this chapter "
  "quantifies the cost of matrix multiplication on a conventional machine, "
  "identifies precisely where that cost is wasted, and shows how the systolic "
  "organisation removes it.")

h2("3.1 Basics of Matrix Multiplication")

p("For the product C = A × B to be defined, the number of columns of A must "
  "equal the number of rows of B. If A is m × k and B is k × n, then C is "
  "m × n and")

equation("Cᵢⱼ = Σ (from p = 0 to k−1) of  Aᵢₚ · Bₚⱼ ,   "
         "0 ≤ i < m,  0 ≤ j < n")

p("Three properties are worth recording. Matrix multiplication is associative, "
  "(AB)C = A(BC), and distributive over addition, A(B + C) = AB + AC, but it "
  "is not commutative: AB ≠ BA in general. The non-commutativity matters in "
  "hardware because it fixes which operand must flow in which direction "
  "through the array; the operand order cannot be swapped for convenience.")

h2("3.2 Normal (Conventional) Matrix Multiplication")

p("The textbook algorithm implements the definition directly as three nested "
  "loops:")

figure(
    "for i = 0 to m-1:              // for each output row\n"
    "    for j = 0 to n-1:          // for each output column\n"
    "        sum = 0\n"
    "        for p = 0 to k-1:      // dot product\n"
    "            sum = sum + A[i][p] * B[p][j]\n"
    "        C[i][j] = sum",
    "The conventional triple-nested-loop matrix multiplication algorithm.")

p("The innermost statement is a multiply-accumulate. The algorithm therefore "
  "performs exactly m·n·k multiplications and m·n·(k−1) additions. For square "
  "matrices of size n this becomes n³ multiplications and n²(n − 1) additions, "
  "giving the familiar O(n³) complexity.")

h2("3.3 Number of Multiplications and Additions in a 4×4 Matrix")

p("Applying these formulas to the 4×4 case that this project implements:")

equation("Number of multiplications = n³ = 4³ = 64")
equation("Number of additions = n²(n − 1) = 16 × 3 = 48")

p("Each of the sixteen output elements requires four multiplications and three "
  "additions. If the accumulator is initialised to zero rather than to the "
  "first product — which is what the hardware in this project does — then each "
  "output element performs four additions, giving 64 multiply-accumulate "
  "operations in total. This figure of 64 MACs is the benchmark against which "
  "the systolic implementation is measured throughout the report.")

table("Operation count for square matrix multiplication.",
      ["Matrix Size (n × n)", "Multiplications (n³)", "Additions (n²(n−1))",
       "Total Operations"],
      [["2 × 2", "8", "4", "12"],
       ["4 × 4", "64", "48", "112"],
       ["8 × 8", "512", "448", "960"],
       ["16 × 16", "4 096", "3 840", "7 936"],
       ["64 × 64", "262 144", "258 048", "520 192"],
       ["256 × 256", "16 777 216", "16 711 680", "33 488 896"]],
      widths=[1.7, 1.7, 1.6, 1.5])

p("The table makes the scaling problem vivid. Doubling the matrix dimension "
  "multiplies the work by eight. A 256×256 product — a modest size by neural "
  "network standards — already requires over sixteen million multiplications. "
  "A scalar processor capable of one MAC per cycle at 1 GHz would need "
  "approximately 16.8 milliseconds for a single such product; a network with "
  "fifty such layers would take almost a second per inference.")

h2("3.4 Drawbacks of Conventional Matrix Multiplication")

p("The problem with the conventional approach is not the arithmetic count, "
  "which is fixed by the definition of the operation. The problem is everything "
  "the processor must do around each arithmetic operation.")

h3("3.4.1 Excessive Memory Traffic")

p("In the naive triple loop, each execution of the inner statement reads "
  "A[i][p] and B[p][j] from memory and reads and writes the accumulator. For a "
  "4×4 product that is 64 reads of A elements and 64 reads of B elements, even "
  "though A and B together contain only 32 distinct values. Every value is "
  "therefore fetched four times. In an n×n product every value is fetched n "
  "times. The arithmetic intensity — the ratio of operations performed to bytes "
  "moved — is disastrously low, and the computation becomes memory-bound rather "
  "than compute-bound.")

h3("3.4.2 The Energy Cost of Data Movement")

p("Memory traffic is not merely slow; it is the dominant consumer of energy. "
  "Measurements widely cited in the computer architecture literature "
  "(Horowitz, ISSCC 2014, for a 45 nm process) give the following approximate "
  "energy costs:")

table("Relative energy cost of arithmetic and data movement (45 nm process).",
      ["Operation", "Approximate Energy", "Relative to 8-bit Add"],
      [["8-bit integer addition", "0.03 pJ", "1×"],
       ["8-bit integer multiplication", "0.20 pJ", "≈ 7×"],
       ["32-bit integer addition", "0.10 pJ", "≈ 3×"],
       ["32-bit SRAM read (8 KB)", "5 pJ", "≈ 170×"],
       ["32-bit DRAM read", "640 pJ", "≈ 21 000×"]],
      widths=[2.2, 1.9, 1.9])

p("A single DRAM access costs thousands of times more energy than the "
  "arithmetic it feeds. It follows that an accelerator is not made efficient by "
  "building faster multipliers; it is made efficient by arranging for each "
  "value fetched from memory to be used as many times as possible before it is "
  "discarded. This is the governing principle of the entire design.")

h3("3.4.3 Instruction Overhead")

p("On a general-purpose processor, each multiply-accumulate is surrounded by "
  "instruction fetch, decode, address generation, loop-counter increment and "
  "branch prediction. The useful arithmetic may represent less than ten per "
  "cent of the energy consumed and only a fraction of the issue slots used. "
  "A systolic array eliminates this overhead entirely: the array executes no "
  "instructions at all, its behaviour being determined by fixed wiring and a "
  "small control state machine.")

h3("3.4.4 Limited Parallelism")

p("Although the m·n dot products are mutually independent and could in "
  "principle all proceed simultaneously, a scalar processor executes them one "
  "at a time. SIMD extensions widen this to a handful of lanes, but each lane "
  "still requires its own operand delivery from the register file, and the "
  "register file's port count becomes the bottleneck. Parallelism in a "
  "general-purpose machine is limited by operand bandwidth, not by the number "
  "of multipliers that could be built.")

h2("3.5 Matrix Multiplication Using a Systolic Array")

p("The systolic array attacks all four drawbacks simultaneously. Consider the "
  "4×4 product C = A × B mapped onto a 4×4 array of processing elements.")

p("The weight matrix B is loaded into the array once, with element B[p][j] "
  "stored permanently in the processing element at row p, column j. The "
  "activation matrix A is then streamed in from the left edge, row p of the "
  "array receiving the p-th element of each row of A. Partial sums enter at "
  "the top of each column, are accumulated as they descend, and emerge "
  "completed at the bottom.")

p("The mapping is exact. Column j of the array computes, for each row i of A, "
  "the sum Σ A[i][p]·B[p][j] — which is precisely C[i][j]. The vertical "
  "position within the column corresponds to the summation index p, so the "
  "reduction is performed by the physical act of the partial sum travelling "
  "down the column. No accumulator is addressed, no loop counter is "
  "incremented, and no intermediate result is written to memory.")

p("The consequences for data movement are dramatic. Each element of B is "
  "fetched from memory exactly once, ever. Each element of A is fetched once "
  "and then reused by all four processing elements in its row as it propagates "
  "rightward. Each partial sum is created, consumed and destroyed inside the "
  "array. The 128 operand fetches of the naive algorithm are reduced to 32.")

h2("3.6 Comparison: Conventional versus Systolic Matrix Multiplication")

table("Conventional processor versus 4×4 systolic array for a 4×4 × 4×4 product.",
      ["Aspect", "Conventional (scalar CPU)", "4×4 Systolic Array"],
      [["MAC operations", "64", "64 (identical)"],
       ["MACs per clock cycle", "1", "16"],
       ["Compute cycles", "64 (plus loop overhead)", "11 (after weight load)"],
       ["Weight fetches from memory", "64", "16 (once each)"],
       ["Activation fetches", "64", "16 (once each)"],
       ["Intermediate results to memory", "48 partial sums", "0"],
       ["Instructions executed", "Hundreds", "0 inside the array"],
       ["Interconnect", "Global bus, register file", "Nearest-neighbour only"],
       ["Scalability", "Limited by operand bandwidth", "Tile-and-abut replication"],
       ["Flexibility", "Any algorithm", "Matrix multiplication only"]],
      widths=[1.9, 2.1, 1.9])

p("The final row is important and is not a defect of the analysis. A systolic "
  "array is not a general-purpose processor and cannot become one. It is a "
  "fixed-function engine that performs one operation extraordinarily well. In "
  "a practical system it is a coprocessor: a host processor handles control "
  "flow, data marshalling and everything that is not a matrix product, and "
  "delegates the matrix product to the array. This is exactly the arrangement "
  "used in Google's Tensor Processing Unit, examined in Chapter 7.")

summary(
    "Conventional matrix multiplication requires n³ multiplications and "
    "n²(n − 1) additions — for a 4×4 product, 64 multiplications and 48 "
    "additions. Its inefficiency lies not in this arithmetic but in the "
    "surrounding overhead: every operand is fetched n times, DRAM accesses cost "
    "thousands of times more energy than the arithmetic they feed, each MAC "
    "carries instruction-processing overhead, and available parallelism goes "
    "unexploited. A systolic array removes all four penalties by loading each "
    "weight once, reusing each activation across an entire row, performing the "
    "summation by physically passing partial sums down a column, and executing "
    "no instructions at all. The next chapter opens up the single processing "
    "element that makes this possible.")

# ===========================================================================
# CHAPTER 4 — PROCESSING ELEMENT
# ===========================================================================

h1("Chapter 4: The Processing Element (PE)")

p("The processing element is the atom of the systolic array. The entire 4×4 "
  "array is nothing more than sixteen copies of the circuit described in this "
  "chapter, wired to their neighbours. Understanding one processing element "
  "completely is therefore equivalent to understanding the whole machine, and "
  "this chapter accordingly examines every block, every signal and every clock "
  "edge in detail.")

h2("4.1 What is a Processing Element?")

p("A processing element is a small synchronous digital circuit that performs "
  "one multiply-accumulate operation per clock cycle and forwards data to its "
  "neighbours. It contains exactly three functional resources: a multiplier, "
  "an adder, and a small set of registers. It contains no instruction memory, "
  "no decoder, no register file and no data memory. Its behaviour is fixed by "
  "its wiring and modulated only by two control signals.")

p("Everything in the design is driven by a single requirement: the processing "
  "element must sustain one MAC per clock cycle indefinitely, without stalling "
  "and without any dependence on how large the surrounding array is. This "
  "forces every interface to be registered and every path between registers to "
  "be short.")

h2("4.2 Internal Block Diagram")

figure(
    "                              wload\n"
    "                                │\n"
    "         psum_in [31:0] ────────┼──────────────────────┬─────────────┐\n"
    "               │                │                      │             │\n"
    "               │ [7:0]          ▼                      │             │\n"
    "               └────────►┌─────────────┐               │             │\n"
    "                         │  weight_q   │               │             │\n"
    "                         │  (8-bit)    │               │             │\n"
    "                         │  REGISTER   │──────┐        │             │\n"
    "                         └─────────────┘      │        │             │\n"
    "                                 │            ▼        ▼             │\n"
    "                                 │        ┌────────────────┐         │\n"
    "         a_in [7:0] ─────┬───────┼───────►│   MULTIPLIER   │         │\n"
    "                         │       │        │   8 × 8 → 16   │         │\n"
    "                         │       │        └───────┬────────┘         │\n"
    "                         │       │                │ sign-extend      │\n"
    "                         │       │                ▼ to 32            │\n"
    "                         │       │        ┌────────────────┐         │\n"
    "                         │       │        │  ADDER 32-bit  │◄────────┤\n"
    "                         │       │        └───────┬────────┘         │\n"
    "                         │       │                ▼                  │\n"
    "                         │       │        ┌───────────────┐          │\n"
    "        valid_in ────────┼───────┼───────►│  MUX (valid)  │◄─────────┘\n"
    "                         │       │        └───────┬───────┘\n"
    "                         ▼       │                ▼\n"
    "                  ┌───────────┐  │        ┌───────────────┐\n"
    "                  │   a_q     │  │        │    psum_q     │\n"
    "                  │ REGISTER  │  │        │  32-bit REG   │\n"
    "                  └─────┬─────┘  │        └───────┬───────┘\n"
    "                        │        │                │\n"
    "                        ▼        │  sign-extended │\n"
    "                     a_out       └───► weight ───►┌──────────┐\n"
    "                    [7:0]                         │   MUX    │◄── wload\n"
    "                                                  └────┬─────┘\n"
    "        valid_in ──►┌──────────┐                       ▼\n"
    "                    │ valid_q  │──► valid_out     psum_out [31:0]\n"
    "                    └──────────┘",
    "Internal block diagram of the weight-stationary processing element, showing "
    "the weight register, the multiplier, the 32-bit accumulating adder, the "
    "activation and valid pipeline registers, and the dual-purpose psum path.")

p("The diagram contains four registers and three combinational blocks. Each is "
  "described in Section 4.5, but the overall shape is worth noting first: "
  "there is exactly one register between any input and any output, so the "
  "processing element adds exactly one cycle of latency to every path through "
  "it, and the longest combinational path is multiplier plus adder.")

h2("4.3 Pin Diagram")

figure(
    "                          psum_in [31:0]\n"
    "                                │\n"
    "                                ▼\n"
    "                  ┌───────────────────────────┐\n"
    "                  │                           │\n"
    "        clk ─────►│                           │\n"
    "                  │                           │\n"
    "      rst_n ─────►│                           │\n"
    "                  │      PROCESSING           │\n"
    "      wload ─────►│        ELEMENT            │\n"
    "                  │                           │\n"
    "   valid_in ─────►│      IN_W  = 8            │────► valid_out\n"
    "                  │      ACC_W = 32           │\n"
    " a_in [7:0] ─────►│                           │────► a_out [7:0]\n"
    "                  │                           │\n"
    "                  └─────────────┬─────────────┘\n"
    "                                │\n"
    "                                ▼\n"
    "                          psum_out [31:0]",
    "Pin diagram of the processing element. Activations enter from the left and "
    "leave to the right; partial sums enter from the top and leave at the "
    "bottom; control enters from the left edge.")

table("Complete port list of the processing element.",
      ["Signal", "Direction", "Width", "Function"],
      [["clk", "Input", "1", "Global clock; all registers update on the rising edge"],
       ["rst_n", "Input", "1", "Active-low synchronous reset; clears all registers"],
       ["wload", "Input", "1", "1 = weight-load phase, 0 = compute phase"],
       ["valid_in", "Input", "1", "1 = a_in carries a valid activation this cycle"],
       ["a_in", "Input", "8", "Incoming activation (signed) from the left neighbour"],
       ["a_out", "Output", "8", "Registered activation forwarded to the right neighbour"],
       ["valid_out", "Output", "1", "Registered valid flag forwarded to the right"],
       ["psum_in", "Input", "32", "Partial sum from above, or weight during wload"],
       ["psum_out", "Output", "32", "Partial sum to below, or weight during wload"]],
      widths=[1.1, 1.0, 0.7, 3.1])

h2("4.4 The MAC Operation")

p("The arithmetic performed by the processing element on each compute cycle is")

equation("psum_out = psum_in + (weight_q × a_in)")

p("Three implementation details make this equation correct in hardware.")

h3("4.4.1 Signed Arithmetic")

p("Both operands are 8-bit two's complement signed values in the range −128 to "
  "+127. The multiplier must therefore be a signed multiplier: both operands "
  "are sign-extended before partial products are formed. Treating a negative "
  "weight as an unsigned value would, for example, interpret −1 as +255 and "
  "corrupt the result completely.")

h3("4.4.2 Product and Accumulator Width")

p("The product of two 8-bit signed values requires 16 bits: the extreme case "
  "(−128) × (−128) = +16 384 does not fit in 15 bits. The product is then "
  "sign-extended to the full 32-bit accumulator width before being added to "
  "the incoming partial sum. The 32-bit accumulator is deliberately generous: "
  "it allows more than 260 000 worst-case products to be summed without "
  "overflow, which is far beyond the four terms required by a 4×4 array and "
  "leaves ample margin for larger arrays built from the same cell.")

equation("Guard bits available = 32 − 16 = 16 bits  ⇒  up to 2¹⁶ = 65 536 "
         "worst-case accumulations")

h3("4.4.3 Conditional Accumulation")

p("The multiply-accumulate is performed only when valid_in is asserted. When "
  "valid_in is low, the incoming partial sum is passed to the output unchanged. "
  "This is essential during pipeline fill and drain, when some processing "
  "elements have no meaningful activation to work with: without this gating, "
  "the array would accumulate the products of whatever residual values happened "
  "to be present on the activation wires, silently corrupting every result.")

equation("psum_q ←  psum_in + weight_q × a_in   if valid_in = 1\n"
         "         psum_in                       if valid_in = 0")

h2("4.5 Detailed Description of Every Block and Signal")

h3("4.5.1 The Weight Register (weight_q)")

p("The weight register is an 8-bit flip-flop bank that stores the stationary "
  "coefficient. It loads a new value from the low eight bits of psum_in on any "
  "rising clock edge on which wload is asserted, and holds its value on every "
  "other edge. Because the register has an explicit enable, it is written only "
  "during the weight-load phase; the compute traffic that passes over the same "
  "psum wires cannot disturb it. This hold behaviour is what makes the "
  "architecture \"weight-stationary\", and it is the single most important "
  "property of the cell.")

p("Only the low eight bits of the 32-bit psum bus are captured. The upper "
  "twenty-four bits are ignored entirely during a weight load, which permits "
  "the surrounding logic to leave them at any convenient value.")

h3("4.5.2 The Activation Register (a_q)")

p("The activation register captures a_in on every rising clock edge, "
  "unconditionally, and drives a_out. Its purpose is to break the horizontal "
  "combinational path: without it, an activation entering the left edge of a "
  "4×4 array would have to propagate through four processing elements within a "
  "single clock period, and the maximum clock frequency would fall as the array "
  "widened. With it, the activation advances exactly one processing element per "
  "cycle regardless of array size — the defining rhythm of a systolic array.")

p("Note carefully that the multiplier is driven by a_in, the incoming value, "
  "not by a_q, the registered value. The processing element multiplies the "
  "activation that is present on its input during the current cycle and "
  "simultaneously registers that same activation for the benefit of its "
  "right-hand neighbour. Multiplying a_q instead would delay the arithmetic by "
  "one cycle relative to the data movement and destroy the alignment between "
  "activations and partial sums.")

h3("4.5.3 The Valid Register (valid_q)")

p("The valid register is a single flip-flop that travels alongside the "
  "activation register, capturing valid_in and driving valid_out. It ensures "
  "that the qualifier arrives at the neighbouring processing element on exactly "
  "the same cycle as the activation it qualifies. A skew of even one cycle "
  "between the two would cause a processing element either to accumulate an "
  "invalid activation or to discard a valid one.")

h3("4.5.4 The Partial-Sum Register (psum_q)")

p("The partial-sum register is a 32-bit flip-flop bank holding the accumulated "
  "value. On every rising clock edge it loads either psum_in + weight_q × a_in "
  "(when valid_in is high) or psum_in unchanged (when valid_in is low). Note "
  "that it never holds its previous value: the accumulation is not local to the "
  "processing element but distributed down the column, with each cell adding "
  "its own contribution to a sum that is passing through. This is the essential "
  "difference between a weight-stationary array and an output-stationary array, "
  "and it is what allows a new independent dot product to enter the column on "
  "every single cycle.")

h3("4.5.5 The Multiplier")

p("The multiplier is an 8 × 8 signed combinational multiplier producing a "
  "16-bit result. In an ASIC or FPGA implementation this is typically mapped to "
  "a dedicated DSP block. Together with the adder it forms the critical path of "
  "the entire array, and therefore determines the maximum clock frequency.")

h3("4.5.6 The Adder")

p("The adder is a 32-bit signed adder that sums the sign-extended product with "
  "psum_in. Because it is 32 bits wide it is a significant contributor to the "
  "critical path; in a high-frequency implementation it would be pipelined or "
  "replaced with a carry-save structure, at the cost of additional latency.")

h3("4.5.7 The Dual-Purpose psum Path")

p("A design decision worth particular attention is that the psum wires carry "
  "two entirely different kinds of information at two different times. During "
  "the weight-load phase (wload = 1) they carry weights being shifted down the "
  "column; during the compute phase (wload = 0) they carry partial sums. The "
  "output multiplexer selects accordingly: it drives the sign-extended contents "
  "of the weight register when wload is high, and the partial-sum register when "
  "wload is low.")

p("The benefit is architectural rather than merely economical. Because the "
  "weight register drives psum_out during a load, each column behaves as a "
  "vertical shift register, and the weights for an entire column can be pushed "
  "in one per cycle from a single port at the top of the array. A 4×4 array "
  "requires only four 32-bit input ports at its top edge to load all sixteen "
  "weights, instead of sixteen separate weight-delivery paths. The saving grows "
  "quadratically with array size: a 256×256 array loads 65 536 weights through "
  "just 256 ports. Without this time-sharing, weight distribution would dominate "
  "the wiring of the chip.")

p("The sign extension in this path is not a detail to be overlooked. A weight "
  "of −5, held as 8'hFB in the weight register, must appear on the 32-bit bus "
  "as 32'hFFFF_FFFB and not as 32'h0000_00FB; otherwise the processing element "
  "below would receive +251 instead of −5.")

h2("4.6 Working of One Processing Element")

p("The complete cycle-by-cycle behaviour of a single processing element is best "
  "seen through a worked example. Assume the weight −3 is to be loaded, after "
  "which the activations 4, 5 and 6 arrive on successive cycles with incoming "
  "partial sums 100, 200 and 300.")

table("Cycle-by-cycle operation of a single processing element.",
      ["Cycle", "wload", "valid_in", "a_in", "psum_in", "weight_q after edge",
       "psum_out after edge"],
      [["0", "1", "0", "—", "−3", "−3", "−3 (weight)"],
       ["1", "0", "1", "4", "100", "−3", "100 + (−3×4) = 88"],
       ["2", "0", "1", "5", "200", "−3", "200 + (−3×5) = 185"],
       ["3", "0", "0", "9", "300", "−3", "300 (unchanged)"],
       ["4", "0", "1", "6", "300", "−3", "300 + (−3×6) = 282"]],
      widths=[0.6, 0.6, 0.7, 0.6, 0.8, 1.2, 1.4])

p("Cycle 0 loads the weight; note that psum_out carries the weight itself, "
  "ready to be captured by the processing element below on the next cycle. "
  "Cycles 1 and 2 perform normal multiply-accumulate operations. Cycle 3 "
  "demonstrates the importance of the valid gate: although a_in carries the "
  "value 9, valid_in is low, so the partial sum passes through untouched. "
  "Cycle 4 resumes normal operation. Throughout, weight_q never changes — the "
  "weight is stationary.")

h2("4.7 Data Movement Between Processing Elements")

p("Within the array, each processing element is connected to at most four "
  "neighbours, and each connection carries a specific stream:")

bullets([
    ("Left neighbour → this PE", "a_in and valid_in arrive from the a_out and "
     "valid_out of the processing element to the left. Processing elements in "
     "column 0 receive these signals from the array boundary."),
    ("This PE → right neighbour", "a_out and valid_out drive the a_in and "
     "valid_in of the processing element to the right. In the rightmost column "
     "these outputs are left unconnected; the activation has been fully "
     "consumed."),
    ("Above → this PE", "psum_in arrives from the psum_out of the processing "
     "element above. Processing elements in row 0 receive either zero (for a "
     "fresh accumulation) or a weight (during the load phase) from the top "
     "boundary."),
    ("This PE → below", "psum_out drives the psum_in of the processing element "
     "below. In the bottom row this output is the completed dot product and "
     "leaves the array."),
])

figure(
    "                       psum_in                  psum_in\n"
    "                          │                        │\n"
    "                          ▼                        ▼\n"
    "     a_in,valid_in  ┌───────────┐  a_out    ┌───────────┐\n"
    "     ──────────────►│  PE(i,j)  ├──────────►│ PE(i,j+1) ├──────────►\n"
    "                    │  W = w₁   │  valid_out│  W = w₂   │\n"
    "                    └─────┬─────┘           └─────┬─────┘\n"
    "                          │ psum_out              │ psum_out\n"
    "                          ▼                       ▼\n"
    "                    ┌───────────┐           ┌───────────┐\n"
    "     ──────────────►│ PE(i+1,j) ├──────────►│PE(i+1,j+1)├──────────►\n"
    "                    │  W = w₃   │           │  W = w₄   │\n"
    "                    └─────┬─────┘           └─────┬─────┘\n"
    "                          ▼                       ▼",
    "Interconnection of neighbouring processing elements. Every connection is "
    "point-to-point and short; there is no global bus anywhere in the array.")

p("Because every one of these connections is registered inside the processing "
  "element, the wire between two adjacent cells is the only combinational path "
  "between them. This is what allows the array to be enlarged without any "
  "reduction in clock frequency, and it is the property that makes systolic "
  "arrays scale to the hundreds-by-hundreds sizes used in commercial "
  "accelerators.")

summary(
    "The processing element is a small synchronous circuit containing one 8×8 "
    "signed multiplier, one 32-bit adder and four registers: weight_q for the "
    "stationary coefficient, a_q and valid_q for the horizontal activation "
    "pipeline, and psum_q for the vertical partial-sum pipeline. It computes "
    "psum_out = psum_in + weight_q × a_in whenever valid_in is asserted and "
    "passes the partial sum through unchanged otherwise. The weight register "
    "loads only when wload is asserted, which is what makes the architecture "
    "weight-stationary; the psum wires are time-shared between weight "
    "distribution and partial-sum accumulation, which reduces the number of "
    "array input ports from n² to n. Because every interface is registered, the "
    "cell adds exactly one cycle of latency in each direction and can be "
    "replicated indefinitely without lengthening the critical path. The next "
    "chapter examines why the weight-stationary choice was made in preference "
    "to the alternatives.")

# ===========================================================================
# CHAPTER 5 — WEIGHT-STATIONARY ARCHITECTURE
# ===========================================================================

h1("Chapter 5: Weight-Stationary Architecture")

p("A systolic array must decide which of the three data streams — weights, "
  "activations or partial sums — is held fixed inside the processing elements "
  "and which are allowed to move. This decision is called the dataflow, and it "
  "is the single most consequential architectural choice in the design of an "
  "accelerator. This chapter explains the weight-stationary dataflow adopted in "
  "this project and justifies it against the alternatives.")

h2("5.1 What is Weight Stationary?")

p("In a weight-stationary dataflow, each processing element holds one weight in "
  "an internal register for the entire duration of a computation. Activations "
  "flow through the array horizontally and partial sums flow vertically, but "
  "the weights never move once loaded.")

p("Operation therefore proceeds in two distinct phases. In the weight-load "
  "phase, which for a 4×4 array occupies four clock cycles, the sixteen weights "
  "are shifted into position through the shared psum wires. In the compute "
  "phase, which may last for thousands of cycles, activations stream through "
  "the array and results emerge continuously. The cost of the load phase is "
  "amortised over the whole compute phase, and the longer the compute phase, "
  "the closer the array comes to its theoretical peak throughput.")

h2("5.2 Why Weights are Stored Inside the PE")

p("The justification comes directly from the structure of neural network "
  "inference. In a trained network the weights are constants: they are computed "
  "once during training and then used unchanged for every subsequent inference. "
  "Activations, by contrast, are different for every input. When a batch of "
  "inputs — or a stream of successive inputs — is processed against the same "
  "layer, the same weight matrix is required over and over while the "
  "activations change continuously.")

p("Placing the invariant operand in a register directly beside the multiplier "
  "that needs it eliminates its memory traffic entirely for the duration of the "
  "computation. The energy table in Section 3.4.2 explains why this is decisive: "
  "reading a weight from DRAM costs on the order of a thousand times more energy "
  "than the multiplication it feeds. If a weight is fetched once and used ten "
  "thousand times, that fetch cost becomes negligible; if it is fetched afresh "
  "for every use, it dominates the entire energy budget of the accelerator.")

p("There is a second, structural reason. The weight register is read by the "
  "multiplier over a wire a few micrometres long, entirely inside the cell. "
  "There is no bus arbitration, no cache lookup, no address decoding and no "
  "possibility of a stall. The operand is simply always there, which is what "
  "permits the guarantee of one MAC per processing element per cycle with no "
  "exceptions.")

h2("5.3 Data Reuse")

p("Data reuse is the quantity that determines the efficiency of any "
  "accelerator. It is defined as the number of arithmetic operations performed "
  "per operand fetched from memory. The weight-stationary array achieves high "
  "reuse on both operands simultaneously, by two different mechanisms.")

h3("5.3.1 Weight Reuse — Temporal")

p("A weight loaded into a processing element is used once per cycle for every "
  "cycle of the compute phase. If M rows of activations are streamed through "
  "the array, each weight is used M times after a single fetch:")

equation("Weight reuse factor = M   (number of activation rows processed)")

p("For a neural network layer processed with a batch of 128 inputs, each weight "
  "is fetched once and used 128 times. This reuse is temporal: the same "
  "physical storage location is read repeatedly over time.")

h3("5.3.2 Activation Reuse — Spatial")

p("An activation entering the left edge of the array propagates rightward "
  "through every processing element in its row. In a 4×4 array it is therefore "
  "used by four different multipliers, contributing to four different output "
  "elements, after a single fetch from memory:")

equation("Activation reuse factor = n   (number of array columns)")

p("This reuse is spatial: the value is used by several physically distinct "
  "processing elements as it travels. Note that this is achieved without any "
  "broadcast wire — the value is relayed from neighbour to neighbour, so the "
  "electrical load on each driver remains that of a single adjacent cell.")

h3("5.3.3 Partial-Sum Reuse — Structural")

p("The third stream is not reused so much as eliminated. A partial sum is "
  "created in one processing element and consumed by the cell immediately "
  "below on the following cycle. It is never written to memory and never read "
  "back. In a 4×4 array, of the four values that make up each output element's "
  "accumulation chain, three exist only inside the array. In a 256×256 array, "
  "255 of every 256 intermediate values never leave the chip.")

table("Data reuse achieved by the weight-stationary 4×4 array.",
      ["Data Type", "Reuse Mechanism", "Reuse Factor", "Memory Accesses"],
      [["Weights", "Temporal (held in register)", "M rows", "16 total, once each"],
       ["Activations", "Spatial (relayed across row)", "n = 4", "One fetch per element"],
       ["Partial sums", "Structural (stay in array)", "—", "Zero"],
       ["Results", "—", "—", "One write per output"]],
      widths=[1.3, 2.2, 1.2, 1.7])

h2("5.4 Broadcast versus Propagation")

p("A naive way to deliver the same activation to four processing elements is "
  "to broadcast it on a shared wire spanning the row. This appears attractive "
  "because it is conceptually simple and introduces no delay, but it scales "
  "badly for three reasons. The capacitive load on the driver grows in "
  "proportion to the number of cells; the wire itself becomes long, and its RC "
  "delay grows with the square of its length; and the resulting long, "
  "high-fanout net becomes the critical path, forcing the clock frequency down "
  "as the array grows.")

p("The systolic array instead uses propagation: the activation is registered in "
  "each cell and handed to the next. The electrical load is constant, the wires "
  "are uniformly short, and the maximum frequency is independent of array size. "
  "The price is latency — the value reaches column j after j cycles — and this "
  "is exactly why the input data must be skewed, as described in Chapter 6. "
  "Propagation converts a physical scaling problem into a scheduling problem, "
  "which is a very favourable exchange.")

table("Broadcast versus systolic propagation.",
      ["Property", "Broadcast", "Propagation (Systolic)"],
      [["Wire length", "Spans entire row", "Cell-to-cell only"],
       ["Driver load", "Grows with n", "Constant (one cell)"],
       ["Critical path", "Degrades with n", "Independent of n"],
       ["Delay to column j", "Zero", "j cycles"],
       ["Scheduling", "Simple", "Requires input skew"],
       ["Scalability", "Poor", "Excellent"]],
      widths=[1.8, 2.1, 2.1])

h2("5.5 Memory Bandwidth")

p("The effect of the weight-stationary dataflow on memory bandwidth is best "
  "seen by direct comparison. Consider processing M rows of activations through "
  "an n × n array, performing M·n² multiply-accumulate operations in total.")

p("A naive implementation fetches two operands per MAC, requiring 2·M·n² "
  "operand fetches. The weight-stationary array fetches n² weights once, and "
  "M·n activation values once each:")

equation("Naive fetches = 2·M·n²")
equation("Weight-stationary fetches = n² + M·n")

p("For M = 64 rows through a 4×4 array, this is 2 048 fetches against 272 — a "
  "reduction of approximately 7.5 times. For a 256×256 array processing 256 "
  "rows, the naive count is 33.5 million and the weight-stationary count is "
  "131 072, a reduction of more than 250 times. The advantage grows with array "
  "size, which is precisely why commercial accelerators use very large arrays.")

p("Expressed differently, the arithmetic intensity — MACs performed per operand "
  "fetched — rises from 0.5 in the naive case to")

equation("Arithmetic intensity = M·n² / (n² + M·n)  →  n  as M grows large")

p("The array's arithmetic intensity therefore approaches n, the dimension of "
  "the array. This single expression captures the entire economic argument for "
  "systolic acceleration.")

h2("5.6 Advantages of the Weight-Stationary Dataflow")

bullets([
    ("Maximum weight reuse", "each weight is fetched once and used for the "
     "entire computation, eliminating the largest source of memory traffic in "
     "neural network inference."),
    ("Very low weight bandwidth", "after the load phase, the weight bandwidth "
     "requirement is exactly zero, which frees the entire memory interface for "
     "activations."),
    ("Simple control", "the array has only two modes, load and compute, "
     "selected by a single control bit. No addressing logic or scheduling "
     "hardware is required inside the array."),
    ("Deterministic timing", "there are no caches, no stalls and no data "
     "dependent delays, so execution time is exactly predictable — a decisive "
     "advantage for real-time and safety-critical systems."),
    ("Short critical path", "the longest combinational path is one multiplier "
     "plus one adder, and it does not lengthen as the array grows."),
    ("Excellent scalability", "the array is built by tiling one cell, so "
     "physical design, verification and timing closure all scale linearly with "
     "area rather than with complexity."),
    ("Ideal for inference", "network weights are constant during inference, "
     "which is exactly the condition under which a stationary weight is most "
     "profitable."),
])

h2("5.7 Disadvantages and Limitations")

p("Intellectual honesty requires an equally careful account of the costs.")

bullets([
    ("Weight-reload overhead", "changing the weight matrix costs n cycles "
     "during which no useful arithmetic occurs. If the array must switch "
     "weights frequently — for instance when each layer is small — this "
     "overhead becomes significant."),
    ("Poor utilisation on small matrices", "a matrix smaller than the array "
     "leaves processing elements idle, and their power is consumed for no "
     "result. A 3×3 layer mapped onto a 256×256 array uses 0.014 % of the "
     "hardware."),
    ("Fixed size", "the array dimensions are fixed at fabrication. Larger "
     "matrices must be decomposed into tiles by software, which adds "
     "complexity and requires storage for intermediate tile results."),
    ("Pipeline fill and drain", "2n − 2 cycles at the beginning and end of "
     "each computation are spent partially idle. For small workloads this "
     "overhead is proportionally large."),
    ("Wide accumulator paths", "the vertical psum wires are 32 bits wide, "
     "against 8 bits for the horizontal activation wires, so the vertical "
     "interconnect dominates the wiring area."),
    ("Single-function hardware", "the array performs matrix multiplication and "
     "nothing else. Every other operation — activation functions, pooling, "
     "normalisation — requires separate hardware."),
    ("Input skew required", "operands must be presented in a staggered pattern, "
     "which requires a skew buffer at the array boundary and complicates the "
     "surrounding data path."),
])

h2("5.8 Weight Stationary versus Output Stationary")

p("In an output-stationary dataflow, each processing element owns one element "
  "of the output matrix and accumulates into a local register. Both weights and "
  "activations flow through the array; the partial sum never moves. Only when "
  "the entire dot product is complete is the accumulator read out.")

p("The trade-off is a direct exchange of one kind of reuse for another. "
  "Output-stationary maximises partial-sum reuse — the accumulator is never "
  "written to memory during accumulation, and the wide 32-bit sum never travels "
  "on a wire — but it requires both operands to be supplied to every cell on "
  "every cycle, doubling operand bandwidth. Weight-stationary maximises weight "
  "reuse at the cost of moving partial sums between cells.")

p("For neural network inference the weight-stationary choice is usually "
  "preferable because weights are the larger and more frequently reused "
  "operand. Output-stationary becomes attractive when the reduction dimension "
  "is very long, since a long accumulation chain in a weight-stationary array "
  "requires either a very tall array or repeated re-entry of partial sums.")

table("Weight-stationary versus output-stationary dataflow.",
      ["Criterion", "Weight Stationary (WS)", "Output Stationary (OS)"],
      [["Held in PE", "Weight", "Partial sum / output"],
       ["Moves horizontally", "Activations", "Activations"],
       ["Moves vertically", "Partial sums", "Weights"],
       ["Maximises reuse of", "Weights", "Partial sums"],
       ["Operand bandwidth", "Low (weights loaded once)", "High (both operands stream)"],
       ["Accumulator location", "Distributed down column", "Local to each PE"],
       ["Psum wire width", "Wide (32-bit vertical)", "None between cells"],
       ["Reload overhead", "n cycles per weight set", "None"],
       ["Best suited to", "Inference, large batches", "Long reduction dimensions"],
       ["Example", "Google TPU v1", "Many academic designs"]],
      widths=[1.6, 2.2, 2.2])

h2("5.9 Weight Stationary versus Input Stationary")

p("In an input-stationary dataflow the activation is the resident operand: each "
  "processing element holds one input value while weights stream through and "
  "partial sums accumulate. This is the mirror image of the weight-stationary "
  "scheme, and it is advantageous only when the activation is the more heavily "
  "reused operand — for example in training, where a single activation is used "
  "against many gradient values, or in convolutions with very large channel "
  "counts and small spatial dimensions.")

p("For inference the arrangement is usually unfavourable, because it requires "
  "the weights to be streamed continuously from memory. Since the weight matrix "
  "is typically far larger than the activation vector, this maximises rather "
  "than minimises memory traffic. It is worth noting that a closely related "
  "scheme, the row-stationary dataflow introduced in the Eyeriss accelerator "
  "(Chen, Emer and Sze, 2016), attempts to balance all three forms of reuse "
  "simultaneously; it achieves better energy efficiency for convolutional "
  "layers at the cost of a considerably more complex control structure and "
  "on-chip network.")

table("Comparison of the three principal dataflows.",
      ["Criterion", "Weight Stationary", "Input Stationary", "Output Stationary"],
      [["Stationary operand", "Weight", "Activation", "Partial sum"],
       ["Streamed operands", "Activation, psum", "Weight, psum", "Weight, activation"],
       ["Weight memory traffic", "Minimum", "Maximum", "Moderate"],
       ["Activation traffic", "Moderate", "Minimum", "Moderate"],
       ["Psum traffic", "Moderate (in array)", "Moderate (in array)", "Minimum"],
       ["Control complexity", "Low", "Low", "Moderate"],
       ["Best for", "Inference", "Training / wide channels", "Long reductions"]],
      widths=[1.5, 1.6, 1.6, 1.5])

summary(
    "The weight-stationary dataflow holds one constant weight in each "
    "processing element for the whole of a computation while activations "
    "propagate horizontally and partial sums accumulate vertically. It is "
    "chosen because neural network weights are invariant during inference and "
    "are the most heavily reused operand, and because fetching a weight from "
    "DRAM costs roughly a thousand times the energy of the multiplication it "
    "feeds. The array achieves temporal reuse of weights, spatial reuse of "
    "activations by nearest-neighbour propagation rather than broadcast, and "
    "complete elimination of partial-sum memory traffic, driving arithmetic "
    "intensity towards n — the array dimension. The costs are an n-cycle weight "
    "reload, poor utilisation on matrices smaller than the array, a fixed array "
    "size, 2n − 2 cycles of pipeline fill and drain, and the need to skew the "
    "input data. Against output-stationary and input-stationary alternatives, "
    "weight-stationary is the correct choice for inference workloads. The next "
    "chapter assembles sixteen such processing elements into the complete 4×4 "
    "array and analyses its behaviour cycle by cycle.")

# ===========================================================================
# CHAPTER 6 — 4x4 ARCHITECTURE
# ===========================================================================

h1("Chapter 6: The 4×4 Systolic Array Architecture")

p("This chapter assembles the complete accelerator. Sixteen instances of the "
  "processing element of Chapter 4 are arranged in a 4×4 mesh, operated "
  "according to the weight-stationary dataflow of Chapter 5, and analysed "
  "cycle by cycle to establish latency, throughput and utilisation.")

h2("6.1 Overall Block Diagram")

figure(
    "                    W col0    W col1    W col2    W col3\n"
    "                psum_in[0] psum_in[1] psum_in[2] psum_in[3]\n"
    "                     │          │          │          │\n"
    "                     ▼          ▼          ▼          ▼\n"
    "   A row0 ──────►┌────────┐ ┌────────┐ ┌────────┐ ┌────────┐\n"
    "   a[0],v[0]     │ PE 0,0 ├►│ PE 0,1 ├►│ PE 0,2 ├►│ PE 0,3 │\n"
    "                 │  B₀₀   │ │  B₀₁   │ │  B₀₂   │ │  B₀₃   │\n"
    "                 └───┬────┘ └───┬────┘ └───┬────┘ └───┬────┘\n"
    "                     ▼          ▼          ▼          ▼\n"
    "   A row1 ──────►┌────────┐ ┌────────┐ ┌────────┐ ┌────────┐\n"
    "   a[1],v[1]     │ PE 1,0 ├►│ PE 1,1 ├►│ PE 1,2 ├►│ PE 1,3 │\n"
    "                 │  B₁₀   │ │  B₁₁   │ │  B₁₂   │ │  B₁₃   │\n"
    "                 └───┬────┘ └───┬────┘ └───┬────┘ └───┬────┘\n"
    "                     ▼          ▼          ▼          ▼\n"
    "   A row2 ──────►┌────────┐ ┌────────┐ ┌────────┐ ┌────────┐\n"
    "   a[2],v[2]     │ PE 2,0 ├►│ PE 2,1 ├►│ PE 2,2 ├►│ PE 2,3 │\n"
    "                 │  B₂₀   │ │  B₂₁   │ │  B₂₂   │ │  B₂₃   │\n"
    "                 └───┬────┘ └───┬────┘ └───┬────┘ └───┬────┘\n"
    "                     ▼          ▼          ▼          ▼\n"
    "   A row3 ──────►┌────────┐ ┌────────┐ ┌────────┐ ┌────────┐\n"
    "   a[3],v[3]     │ PE 3,0 ├►│ PE 3,1 ├►│ PE 3,2 ├►│ PE 3,3 │\n"
    "                 │  B₃₀   │ │  B₃₁   │ │  B₃₂   │ │  B₃₃   │\n"
    "                 └───┬────┘ └───┬────┘ └───┬────┘ └───┬────┘\n"
    "                     ▼          ▼          ▼          ▼\n"
    "                  C[·][0]    C[·][1]    C[·][2]    C[·][3]",
    "Overall block diagram of the 4×4 weight-stationary systolic array. "
    "Processing element (p, j) permanently holds weight B[p][j]. Activations "
    "enter from the left, partial sums descend each column, and completed "
    "output elements emerge at the bottom.")

p("The mapping is direct and worth stating precisely. The array computes "
  "C = A × B, where B is the 4×4 weight matrix resident in the array and A is "
  "a matrix of activations streamed in from the left. Processing element (p, j) "
  "holds B[p][j]. Row p of the array is fed the p-th element of each activation "
  "row, and column j of the array produces column j of the result.")

p("The correctness of this mapping follows from the definition of the matrix "
  "product. As the partial sum for output row i descends column j, it visits "
  "processing elements (0, j) through (3, j) in turn, and at each one it "
  "accumulates A[i][p] × B[p][j]. On leaving the bottom of the column it "
  "therefore carries")

equation("Σ (from p = 0 to 3) of  A[i][p] · B[p][j]  =  C[i][j]")

p("The physical descent through the column performs the summation. There is no "
  "explicit reduction step anywhere in the design.")

h2("6.2 The Sixteen Processing Elements")

p("All sixteen processing elements are identical instances of the same module. "
  "They differ only in the weight they hold and in what their boundary "
  "connections are attached to:")

bullets([
    ("Column 0 (left edge)", "a_in and valid_in come from the array's "
     "activation input ports rather than from a neighbour."),
    ("Column 3 (right edge)", "a_out and valid_out are left unconnected; the "
     "activation has been used by every column and is discarded."),
    ("Row 0 (top edge)", "psum_in comes from the array's top-edge input ports: "
     "weights during the load phase, and zero during the compute phase."),
    ("Row 3 (bottom edge)", "psum_out carries completed output elements out of "
     "the array."),
])

p("The wload control signal is broadcast to all sixteen processing elements "
  "simultaneously, since the entire array switches between load and compute "
  "phases as a unit. It is the only global signal in the design apart from the "
  "clock and reset.")

table("Weight assignment and boundary connections of the sixteen processing elements.",
      ["PE", "Holds", "a_in from", "psum_in from", "psum_out to"],
      [["(0,0)…(0,3)", "B[0][0…3]", "Left edge / PE(0,j−1)", "Top edge port", "PE(1,j)"],
       ["(1,0)…(1,3)", "B[1][0…3]", "Left edge / PE(1,j−1)", "PE(0,j)", "PE(2,j)"],
       ["(2,0)…(2,3)", "B[2][0…3]", "Left edge / PE(2,j−1)", "PE(1,j)", "PE(3,j)"],
       ["(3,0)…(3,3)", "B[3][0…3]", "Left edge / PE(3,j−1)", "PE(2,j)", "Array output"]],
      widths=[1.2, 1.2, 1.7, 1.3, 1.2])

h2("6.3 Input and Output Data Flow")

h3("6.3.1 The Weight-Load Phase")

p("Before computation begins, wload is asserted and the sixteen weights are "
  "shifted into the array through the four top-edge psum ports. Because each "
  "processing element drives its own weight register onto psum_out while wload "
  "is high, every column behaves as a four-deep shift register.")

p("The order of loading follows directly from the shift-register behaviour: the "
  "weight that must end up deepest in the column is presented first. Column j "
  "therefore receives B[3][j], then B[2][j], then B[1][j], then B[0][j] over "
  "four consecutive cycles. All four columns are loaded in parallel, so the "
  "entire weight matrix is in place after exactly four cycles.")

table("Weight-load sequence for the 4×4 array (four cycles, all columns in parallel).",
      ["Load Cycle", "Column 0 port", "Column 1 port", "Column 2 port", "Column 3 port"],
      [["L0", "B[3][0]", "B[3][1]", "B[3][2]", "B[3][3]"],
       ["L1", "B[2][0]", "B[2][1]", "B[2][2]", "B[2][3]"],
       ["L2", "B[1][0]", "B[1][1]", "B[1][2]", "B[1][3]"],
       ["L3", "B[0][0]", "B[0][1]", "B[0][2]", "B[0][3]"]],
      widths=[1.1, 1.4, 1.4, 1.4, 1.4])

h3("6.3.2 Input Skew")

p("During the compute phase the activations cannot be presented to the four "
  "array rows simultaneously. An activation entering row p must meet the "
  "descending partial sum for the correct output row at processing element "
  "(p, j), and since the partial sum takes one cycle to move from row p − 1 to "
  "row p, the activation for row p must arrive one cycle later than the "
  "activation for row p − 1. The input data must therefore be skewed: row p of "
  "the array is delayed by p cycles.")

p("Formally, at cycle t, array row p is presented with activation element "
  "A[t − p][p], which is valid whenever 0 ≤ t − p ≤ M − 1. The skew is "
  "implemented by a triangular bank of shift registers at the left edge of the "
  "array — zero registers for row 0, one for row 1, two for row 2 and three for "
  "row 3.")

figure(
    "  cycle:        0     1     2     3     4     5     6\n"
    "               ───────────────────────────────────────►\n"
    "  row 0 →     A₀₀   A₁₀   A₂₀   A₃₀    ·     ·     ·\n"
    "  row 1 →      ·    A₀₁   A₁₁   A₂₁   A₃₁    ·     ·\n"
    "  row 2 →      ·     ·    A₀₂   A₁₂   A₂₂   A₃₂    ·\n"
    "  row 3 →      ·     ·     ·    A₀₃   A₁₃   A₂₃   A₃₃\n"
    "\n"
    "        (Aᵢₚ = element p of activation row i; row p delayed by p cycles)",
    "Skewed activation input schedule. Each successive array row is delayed by "
    "one additional cycle so that activations meet the correct descending "
    "partial sum.")

h3("6.3.3 Output De-skew")

p("The outputs emerge with a complementary skew. Because an activation reaches "
  "column j only after j additional cycles of propagation, the partial sum for "
  "output row i must enter column j at cycle i + j, and the completed result "
  "emerges from the bottom of column j at cycle i + j + 4. Column 0 therefore "
  "produces its results first and column 3 last, three cycles later. A matching "
  "triangular bank of shift registers at the bottom edge re-aligns the four "
  "columns so that a complete output row is presented to the outside world "
  "simultaneously.")

h2("6.4 Clock-by-Clock Operation")

p("The complete operation of the array for the product of a 4×4 activation "
  "matrix with the resident 4×4 weight matrix is set out below. Cycles L0–L3 "
  "are the weight-load phase; cycles 0–10 are the compute phase.")

table("Complete clock-by-clock operation of the 4×4 array.",
      ["Cycle", "Activations entering rows 0–3", "Array activity",
       "Outputs emerging"],
      [["L0–L3", "—", "Weights shifted down all four columns", "—"],
       ["0", "A₀₀ / — / — / —", "PE(0,0) active; pipeline filling", "—"],
       ["1", "A₁₀ / A₀₁ / — / —", "3 PEs active", "—"],
       ["2", "A₂₀ / A₁₁ / A₀₂ / —", "6 PEs active", "—"],
       ["3", "A₃₀ / A₂₁ / A₁₂ / A₀₃", "10 PEs active", "—"],
       ["4", "— / A₃₁ / A₂₂ / A₁₃", "All 16 PEs active (array full)", "C₀₀"],
       ["5", "— / — / A₃₂ / A₂₃", "Draining begins", "C₁₀, C₀₁"],
       ["6", "— / — / — / A₃₃", "Draining", "C₂₀, C₁₁, C₀₂"],
       ["7", "—", "Draining", "C₃₀, C₂₁, C₁₂, C₀₃"],
       ["8", "—", "Draining", "C₃₁, C₂₂, C₁₃"],
       ["9", "—", "Draining", "C₃₂, C₂₃"],
       ["10", "—", "Final result leaves the array", "C₃₃"]],
      widths=[0.8, 2.0, 2.1, 1.6])

p("Reading the table row by row makes the behaviour of the machine concrete. "
  "For the first three cycles the array is filling: activations have entered "
  "but have not yet reached the rightmost columns, and no partial sum has yet "
  "traversed all four rows. At cycle 4 the array is completely full — all "
  "sixteen processing elements perform a useful multiply-accumulate on the same "
  "cycle — and the first result, C₀₀, emerges. From cycle 5 onwards no new "
  "activations enter and the array drains, producing the remaining fifteen "
  "results over the following six cycles. The final result C₃₃ leaves at cycle "
  "10.")

h2("6.5 Pipeline Filling and Draining")

p("The fill and drain periods are an inherent consequence of the propagation "
  "delays that make the array scalable. Their duration follows from the "
  "geometry of the mesh.")

p("Filling: the first activation must travel from the left edge to the "
  "rightmost column (n − 1 cycles of horizontal propagation) and the first "
  "partial sum must travel from the top row to the bottom row (n − 1 cycles of "
  "vertical propagation, plus the final register stage). The array is fully "
  "occupied only from cycle n − 1 onwards.")

p("Draining: after the last activation enters, the results already in flight "
  "must complete their journey through the array, requiring a further "
  "n − 1 cycles of horizontal travel and n cycles of vertical descent.")

equation("Fill time = n − 1 = 3 cycles")
equation("Drain time = n − 1 = 3 cycles")
equation("Total overhead = 2(n − 1) = 6 cycles")

p("This overhead is fixed and independent of how much data is processed. It is "
  "therefore negligible for long computations and significant only for short "
  "ones — which is precisely the reason a weight-stationary array should be "
  "given as many activation rows as possible per weight load.")

h2("6.6 Number of Clock Cycles")

p("A general expression for the total execution time can now be derived. Let "
  "the array be n × n and let M rows of activations be streamed through it. "
  "The last activation row enters array row n − 1 at cycle (M − 1) + (n − 1). "
  "Its contribution to column n − 1 requires a further n − 1 cycles of "
  "horizontal propagation, and the resulting partial sum requires n cycles to "
  "descend the column and be registered at the output. The last result "
  "therefore emerges at cycle")

equation("t_last = (M − 1) + (n − 1) + n  =  M + 2n − 2")

p("Counting from cycle 0, the compute phase occupies")

equation("N_compute = M + 2n − 1  cycles")

p("and the complete operation including the weight load occupies")

equation("N_total = n + M + 2n − 1 = M + 3n − 1  cycles")

p("For the 4×4 × 4×4 case with M = n = 4:")

equation("N_compute = 4 + 8 − 1 = 11 cycles;   N_total = 4 + 12 − 1 = 15 cycles")

table("Execution time of the 4×4 array for different numbers of activation rows.",
      ["Activation rows M", "MAC operations", "Compute cycles (M + 7)",
       "MACs per cycle", "Utilisation"],
      [["4", "64", "11", "5.8", "36 %"],
       ["8", "128", "15", "8.5", "53 %"],
       ["16", "256", "23", "11.1", "70 %"],
       ["64", "1 024", "71", "14.4", "90 %"],
       ["256", "4 096", "263", "15.6", "97 %"],
       ["1 024", "16 384", "1 031", "15.9", "99 %"]],
      widths=[1.3, 1.3, 1.5, 1.2, 1.1])

p("The table demonstrates the amortisation principle quantitatively. "
  "Utilisation, defined as M / (M + 2n − 1), rises steadily towards 100 % as "
  "the number of activation rows grows. A systolic array is at its worst on "
  "tiny problems and at its best on large, sustained streams of data — which "
  "matches the characteristics of batched neural network inference very well.")

h2("6.7 Timing Diagram")

figure(
    "  clk    ┌─┐ ┌─┐ ┌─┐ ┌─┐ ┌─┐ ┌─┐ ┌─┐ ┌─┐ ┌─┐ ┌─┐ ┌─┐ ┌─┐ ┌─┐ ┌─┐ ┌─┐\n"
    "       ──┘ └─┘ └─┘ └─┘ └─┘ └─┘ └─┘ └─┘ └─┘ └─┘ └─┘ └─┘ └─┘ └─┘ └─┘ └──\n"
    "  cycle   L0  L1  L2  L3   0   1   2   3   4   5   6   7   8   9  10\n"
    "\n"
    "  rst_n ──────────────────────────────────────────────────────────────\n"
    "\n"
    "  wload ┌───────────────┐\n"
    "       ─┘               └──────────────────────────────────────────────\n"
    "\n"
    "  psum_in  ╳W3 ╳W2 ╳W1 ╳W0 ╳═══════════ 0 (fresh accumulation) ════════\n"
    "  (top)\n"
    "\n"
    "  valid                  ┌───────────────────────────┐\n"
    "       ────────────────  ┘                           └─────────────────\n"
    "\n"
    "  a_in row0 ─────────────╳A₀₀╳A₁₀╳A₂₀╳A₃₀╳────────────────────────────\n"
    "  a_in row1 ─────────────────╳A₀₁╳A₁₁╳A₂₁╳A₃₁╳────────────────────────\n"
    "  a_in row2 ─────────────────────╳A₀₂╳A₁₂╳A₂₂╳A₃₂╳────────────────────\n"
    "  a_in row3 ─────────────────────────╳A₀₃╳A₁₃╳A₂₃╳A₃₃╳────────────────\n"
    "\n"
    "  psum_out col0 ─────────────────────────╳C₀₀╳C₁₀╳C₂₀╳C₃₀╳────────────\n"
    "  psum_out col1 ─────────────────────────────╳C₀₁╳C₁₁╳C₂₁╳C₃₁╳────────\n"
    "  psum_out col2 ─────────────────────────────────╳C₀₂╳C₁₂╳C₂₂╳C₃₂╳────\n"
    "  psum_out col3 ─────────────────────────────────────╳C₀₃╳C₁₃╳C₂₃╳C₃₃╳\n"
    "\n"
    "            │◄── load ──►│◄─ fill ─►│◄ full ►│◄──── drain ────►│",
    "Timing diagram of the complete 4×4 matrix multiplication, showing the "
    "four-cycle weight load, the skewed activation inputs, the fill and drain "
    "periods, and the column-skewed outputs.")

h2("6.8 Latency and Throughput")

h3("6.8.1 Latency")

p("Latency is the delay from the presentation of the first activation to the "
  "availability of the first result. From Section 6.4, C₀₀ emerges at cycle 4, "
  "so")

equation("Latency = n = 4 cycles  (compute phase only)")
equation("Latency including weight load = 2n = 8 cycles")

p("At a 200 MHz clock this is 20 ns and 40 ns respectively. Latency is a "
  "property of the depth of the array and cannot be reduced without making the "
  "array shallower.")

h3("6.8.2 Throughput")

p("Throughput is the rate at which results are produced once the pipeline is "
  "full. In steady state each of the four columns emits one completed output "
  "element per cycle, so")

equation("Throughput = n = 4 output elements per cycle")

p("Each output element is a four-term dot product, so the sustained arithmetic "
  "rate is")

equation("Peak performance = n² = 16 MAC/cycle = 32 operations/cycle")

p("At 200 MHz this corresponds to 6.4 giga-operations per second from sixteen "
  "small multipliers — a figure a scalar processor would need to run at "
  "3.2 GHz to match, and then only with perfect instruction scheduling and no "
  "memory stalls.")

table("Performance summary of the 4×4 weight-stationary systolic array.",
      ["Metric", "Value", "Comment"],
      [["Processing elements", "16", "4 × 4 mesh"],
       ["Peak throughput", "16 MAC/cycle", "One per PE"],
       ["Operations per cycle", "32", "16 multiplies + 16 adds"],
       ["Latency (first result)", "4 cycles", "Array depth"],
       ["Weight-load time", "4 cycles", "One per array row"],
       ["4×4 × 4×4 product", "11 cycles", "Plus 4-cycle load"],
       ["Scalar CPU equivalent", "64 cycles", "One MAC per cycle"],
       ["Speed-up (compute only)", "≈ 5.8×", "64 / 11"],
       ["Speed-up (sustained)", "16×", "Large M"],
       ["Critical path", "1 multiply + 1 add", "Independent of array size"]],
      widths=[2.0, 1.5, 2.5])

h3("6.8.3 Scaling to Larger Matrices")

p("A matrix larger than 4×4 is handled by tiling. A 16×16 product is "
  "decomposed into sixteen 4×4 sub-products; the array is loaded with each "
  "weight tile in turn and the activation tiles are streamed through, with "
  "partial results accumulated externally. The cost of this decomposition is "
  "one weight reload per tile, which is why commercial designs use much larger "
  "arrays: the 256×256 array of the Tensor Processing Unit handles a 256×256 "
  "layer without any tiling at all.")

summary(
    "The 4×4 array consists of sixteen identical processing elements in which "
    "PE(p, j) holds weight B[p][j]; activations enter from the left with a "
    "one-cycle-per-row skew, partial sums descend each column performing the "
    "summation physically, and results emerge from the bottom with a "
    "complementary column skew. Weights are loaded in four cycles through the "
    "shared psum ports, after which a 4×4 × 4×4 product completes in 11 cycles "
    "against 64 for a scalar processor. Fill and drain consume a fixed "
    "2(n − 1) = 6 cycles, so utilisation rises from 36 % for a single 4×4 "
    "product to 99 % when 1 024 activation rows are streamed against one weight "
    "load. Steady-state throughput is 16 MAC per cycle — one per processing "
    "element — with a latency of 4 cycles and a critical path of one multiply "
    "plus one add that does not lengthen as the array grows. The next chapter "
    "shows these same principles applied at a scale 4 096 times larger in a "
    "commercial product.")

# ===========================================================================
# CHAPTER 7 — GOOGLE TPU
# ===========================================================================

h1("Chapter 7: The Google Tensor Processing Unit")

p("The architecture developed in the preceding chapters is not a purely "
  "academic construction. It is the organising principle of one of the most "
  "commercially significant processors of the last decade. This chapter "
  "examines Google's first-generation Tensor Processing Unit, which is a "
  "weight-stationary systolic array of exactly the kind described in this "
  "report, scaled up by a factor of 4 096 in processing elements.")

h2("7.1 What is a TPU?")

p("A Tensor Processing Unit is an application-specific integrated circuit "
  "designed by Google to accelerate the inference phase of neural networks. "
  "The first generation, described publicly by Jouppi and colleagues at the "
  "International Symposium on Computer Architecture in 2017, had been deployed "
  "in Google's data centres since 2015. Its origin is instructive: Google "
  "projected that if users spoke to Android voice search for three minutes per "
  "day, serving the resulting neural network inference on conventional "
  "processors would require a doubling of the company's entire data centre "
  "capacity. Building a specialised processor was cheaper than building the "
  "data centres.")

p("The TPU is a coprocessor rather than a stand-alone processor. It occupies a "
  "PCIe slot, has no program counter and executes no control flow of its own; a "
  "host CPU sends it high-level instructions such as \"multiply this activation "
  "matrix by that weight matrix\" over the PCIe bus. This division of labour is "
  "exactly the one described at the end of Chapter 3.")

h2("7.2 TPU Architecture")

figure(
    "        ┌──────────────────────────────────────────────────┐\n"
    "        │                  DDR3 DRAM                       │\n"
    "        │              (8 GiB weight memory)               │\n"
    "        └────────────────────┬─────────────────────────────┘\n"
    "                             │  30 GiB/s\n"
    "                             ▼\n"
    "        ┌──────────────────────────────────────────────────┐\n"
    "        │              Weight FIFO (4 deep)                │\n"
    "        └────────────────────┬─────────────────────────────┘\n"
    "                             │\n"
    "   ┌──────────────┐          ▼\n"
    "   │   Unified    │   ┌──────────────────────────────┐\n"
    "   │   Buffer     │──►│    MATRIX MULTIPLY UNIT      │\n"
    "   │  (24 MiB     │   │      256 × 256 systolic      │\n"
    "   │  activation  │   │   array = 65 536 MAC units   │\n"
    "   │   storage)   │   │   (weight-stationary)        │\n"
    "   └──────▲───────┘   └──────────────┬───────────────┘\n"
    "          │                          │\n"
    "          │                          ▼\n"
    "          │           ┌──────────────────────────────┐\n"
    "          │           │   Accumulators (4 MiB)       │\n"
    "          │           └──────────────┬───────────────┘\n"
    "          │                          ▼\n"
    "          │           ┌──────────────────────────────┐\n"
    "          └───────────┤  Activation / Normalise/Pool │\n"
    "                      └──────────────────────────────┘\n"
    "                             ▲\n"
    "                             │ PCIe  (host instructions)\n"
    "                        ┌────┴────┐\n"
    "                        │  HOST   │\n"
    "                        │   CPU   │\n"
    "                        └─────────┘",
    "Simplified block diagram of the first-generation Google Tensor Processing "
    "Unit. The Matrix Multiply Unit is a 256×256 weight-stationary systolic "
    "array containing 65 536 multiply-accumulate units.")

p("The principal blocks are as follows.")

bullets([
    ("Matrix Multiply Unit (MXU)", "a 256 × 256 systolic array of 8-bit "
     "multiply-accumulate units — 65 536 processing elements in total, each "
     "functionally equivalent to the cell designed in Chapter 4."),
    ("Unified Buffer", "24 MiB of on-chip SRAM holding activations. Its size is "
     "deliberate: keeping activations on chip avoids the DRAM accesses that "
     "would otherwise dominate the energy budget."),
    ("Accumulators", "4 MiB of storage for 32-bit partial sums emerging from "
     "the bottom of the array, allowing results from successive passes to be "
     "combined."),
    ("Weight FIFO", "a four-deep buffer that fetches the next weight tile from "
     "DRAM while the current tile is still in use, hiding the weight-load "
     "latency behind computation."),
    ("Activation pipeline", "hardware for the non-matrix operations — ReLU, "
     "normalisation and pooling — that must be applied between layers."),
])

h2("7.3 The Systolic Array Inside the TPU")

p("The Matrix Multiply Unit operates on exactly the principles described in "
  "Chapters 4 to 6. Weights are loaded into the array and held stationary; "
  "activations are read from the Unified Buffer and flow through the array; "
  "partial sums descend the columns and emerge as completed dot products into "
  "the accumulators. Each processing element performs an 8-bit multiply and a "
  "wide accumulate, and each communicates only with its immediate neighbours.")

p("The performance follows from simple multiplication. With 65 536 processing "
  "elements clocked at 700 MHz, each performing one multiply and one add per "
  "cycle:")

equation("65 536 × 2 × 700 × 10⁶ ≈ 92 × 10¹² operations/second = 92 TOPS")

p("This is achieved within a thermal design power of 75 W. The comparison with "
  "a general-purpose processor is stark not because the TPU's transistors are "
  "faster, but because almost all of them are doing arithmetic almost all of "
  "the time. Jouppi and colleagues report that on Google's production inference "
  "workloads the TPU was on average of the order of fifteen to thirty times "
  "faster than the contemporary server CPU and GPU it was compared against, and "
  "delivered roughly thirty to eighty times better performance per watt.")

table("Key parameters of the first-generation Google TPU.",
      ["Parameter", "Value"],
      [["Systolic array size", "256 × 256"],
       ["Multiply-accumulate units", "65 536"],
       ["Arithmetic precision", "8-bit integer multiply, 32-bit accumulate"],
       ["Clock frequency", "700 MHz"],
       ["Peak throughput", "92 TOPS (8-bit integer)"],
       ["On-chip activation memory", "24 MiB Unified Buffer"],
       ["Accumulator memory", "4 MiB"],
       ["External memory", "8 GiB DDR3, ≈ 30 GiB/s"],
       ["Thermal design power", "75 W"],
       ["Dataflow", "Weight stationary"],
       ["Host interface", "PCIe Gen3 ×16"],
       ["Deployment", "Google data centres from 2015"]],
      widths=[2.6, 3.3])

h2("7.4 Why Google Uses Weight Stationary")

p("The reasoning behind Google's choice mirrors the argument of Chapter 5, "
  "applied at data-centre scale.")

bullets([
    ("Inference weights are constant", "a deployed model's weights do not "
     "change between requests, so a weight loaded into the array remains valid "
     "for millions of inferences."),
    ("Weights are the dominant data volume", "a production model may contain "
     "hundreds of millions of parameters against a few thousand activation "
     "values per request; minimising weight traffic minimises total traffic."),
    ("DRAM bandwidth was the binding constraint", "with only about 30 GiB/s of "
     "DDR3 bandwidth, the TPU could not have sustained 92 TOPS if it had to "
     "re-fetch operands. The weight-stationary dataflow reduces the required "
     "bandwidth by orders of magnitude."),
    ("Energy is the true currency of a data centre", "at 75 W the TPU is "
     "limited by what a server slot can cool. Every DRAM access avoided is "
     "energy available for arithmetic."),
    ("Determinism matters for service latency", "user-facing services are "
     "governed by tail latency. A systolic array has no caches and no branch "
     "prediction, so its execution time is exactly predictable — a property "
     "Jouppi and colleagues cite explicitly as a design goal."),
    ("Simplicity permitted rapid development", "the regularity of a systolic "
     "array meant that a single processing element could be designed, verified "
     "and then replicated 65 536 times; the TPU went from concept to deployment "
     "in about fifteen months."),
])

h2("7.5 CPU versus GPU versus TPU")

p("The three classes of processor represent three different answers to the "
  "question of how transistors should be spent.")

h3("7.5.1 The CPU: Optimised for Latency and Generality")

p("A general-purpose CPU devotes the majority of its area to caches, branch "
  "predictors, out-of-order issue logic and register renaming — machinery whose "
  "purpose is to execute an arbitrary, unpredictable instruction stream as "
  "quickly as possible. Only a small fraction of the die performs arithmetic. "
  "This is the correct design for code with complex control flow, but for a "
  "dense matrix product almost all of that machinery is overhead.")

h3("7.5.2 The GPU: Optimised for Throughput")

p("A GPU spends far more of its area on arithmetic units and hides memory "
  "latency by interleaving thousands of threads rather than by predicting "
  "branches. It is enormously more efficient than a CPU for regular parallel "
  "work. It nevertheless retains a general-purpose programming model: every "
  "thread fetches its own operands through a register file and a cache "
  "hierarchy, and that operand-delivery path remains a substantial fraction of "
  "the energy cost.")

h3("7.5.3 The TPU: Optimised for One Operation")

p("The TPU abandons generality altogether. It has no caches, no branch "
  "prediction, no out-of-order execution and no threads. Operands are delivered "
  "not through a register file but by the physical structure of the systolic "
  "array, in which each value is handed directly from one arithmetic unit to "
  "the next. This is why it can devote so much of its area and power budget to "
  "multiply-accumulate units, and why it cannot run anything other than the "
  "operations it was built for.")

table("Architectural comparison of CPU, GPU and TPU.",
      ["Characteristic", "CPU", "GPU", "TPU (v1)"],
      [["Design goal", "Low latency, generality", "High throughput", "One operation, efficiently"],
       ["Arithmetic units", "Tens", "Thousands", "65 536 MACs"],
       ["Control logic", "Very large fraction", "Moderate", "Minimal"],
       ["Cache hierarchy", "Deep (L1/L2/L3)", "Moderate + shared memory", "None (scratchpad only)"],
       ["Operand delivery", "Register file", "Register file per thread", "Neighbour-to-neighbour"],
       ["Programming model", "Sequential threads", "SIMT, thousands of threads", "Host-issued matrix ops"],
       ["Execution timing", "Unpredictable", "Fairly predictable", "Exactly deterministic"],
       ["Precision", "32/64-bit float", "16/32-bit float", "8-bit int, 32-bit accum."],
       ["Flexibility", "Any program", "Any parallel program", "Matrix multiply only"],
       ["Efficiency on MatMul", "Low", "Good", "Very high"]],
      widths=[1.5, 1.4, 1.6, 1.6])

p("The comparison should not be read as a ranking. Each processor is close to "
  "optimal for the workload it was designed for, and a practical system uses "
  "all three: the CPU runs the operating system and application logic, the GPU "
  "handles training and irregular parallel work, and the TPU executes the dense "
  "matrix products of inference. The lesson is that specialisation buys "
  "efficiency at the direct cost of flexibility, and the art of computer "
  "architecture lies in choosing where on that spectrum to sit.")

h2("7.6 Later Generations")

p("Subsequent TPU generations retain the systolic array as their core while "
  "extending its capabilities. The second and third generations added "
  "floating-point support so that training as well as inference could be "
  "accelerated, replaced DDR3 with High Bandwidth Memory, and introduced "
  "multiple matrix units per chip together with a dedicated inter-chip "
  "interconnect for building large pods. Later generations have increased "
  "capacity and interconnect performance further still. Through all of these "
  "changes the fundamental organisation examined in this report — a mesh of "
  "multiply-accumulate cells with stationary weights and neighbour-to-neighbour "
  "communication — has remained unchanged, which is a strong indication of how "
  "well matched it is to the problem.")

summary(
    "Google's first-generation Tensor Processing Unit is a direct industrial "
    "realisation of the architecture developed in this report: a 256 × 256 "
    "weight-stationary systolic array of 65 536 8-bit multiply-accumulate units "
    "operating at 700 MHz to deliver 92 TOPS within a 75 W power budget, "
    "supported by a 24 MiB on-chip Unified Buffer and driven by a host CPU over "
    "PCIe. Google chose the weight-stationary dataflow because inference "
    "weights are constant, weights dominate data volume, DRAM bandwidth was the "
    "binding constraint, energy is the currency of the data centre, and "
    "deterministic timing is essential for user-facing latency guarantees. "
    "Compared with a CPU optimised for latency and generality and a GPU "
    "optimised for throughput with a general programming model, the TPU spends "
    "almost all of its transistors on arithmetic by eliminating caches, "
    "prediction and register files in favour of operand delivery through the "
    "structure of the array itself — buying very high efficiency at the price "
    "of doing only one thing.")

# ===========================================================================
# CHAPTER 8 — CONCLUSION
# ===========================================================================

h1("Chapter 8: Conclusion and Future Scope")

h2("8.1 Conclusion")

p("This report has followed a single line of reasoning from the definition of "
  "multiplication to a commercial data-centre accelerator. The argument may be "
  "summarised as follows.")

p("Multiplication implemented in hardware is expensive in area and energy, and "
  "produces results wider than its operands. Matrix multiplication is built "
  "from an enormous number of such multiplications arranged as "
  "multiply-accumulate operations, and it dominates the computational cost of "
  "artificial intelligence, signal processing, graphics, robotics and "
  "scientific computing. Executed conventionally it requires n³ multiplications "
  "and, far more damagingly, re-fetches every operand n times from a memory "
  "system in which a single access costs thousands of times the energy of the "
  "arithmetic it feeds.")

p("The systolic array eliminates this waste by exploiting the reuse inherent in "
  "the operation. Data is pumped rhythmically through a regular mesh of simple "
  "processing elements that communicate only with their neighbours; each "
  "operand entering the array is used many times before it leaves, and "
  "intermediate results never return to memory. The weight-stationary variant "
  "holds the invariant operand of neural network inference — the weight — in a "
  "register beside the multiplier that needs it, reducing weight memory traffic "
  "to a single fetch per processing element per computation.")

p("The processing element designed in this project realises this scheme in a "
  "compact synchronous cell containing an 8-bit signed multiplier, a 32-bit "
  "accumulating adder and four registers, computing "
  "psum_out = psum_in + weight × a_in once per clock cycle. Its psum interface "
  "is deliberately time-shared between weight distribution and partial-sum "
  "accumulation, which reduces the number of array input ports from n² to n and "
  "is the key to loading a large array through a small boundary.")

p("Assembled into a 4×4 mesh, sixteen such cells complete a 4×4 × 4×4 matrix "
  "product in 11 clock cycles after a 4-cycle weight load, against 64 cycles "
  "for a scalar processor. In sustained operation the array delivers 16 "
  "multiply-accumulate operations per cycle with a latency of only 4 cycles and "
  "a critical path of one multiply plus one add that does not lengthen as the "
  "array grows — the property that permits the same design to be scaled to the "
  "256 × 256 array of Google's Tensor Processing Unit, where 65 536 identical "
  "cells deliver 92 TOPS within 75 W.")

p("The wider lesson of the project is architectural rather than arithmetical. "
  "The systolic array is not faster because its multipliers are better; its "
  "multipliers are entirely ordinary. It is faster and vastly more efficient "
  "because it moves data well. In an era in which transistors are abundant and "
  "the energy of communication dominates the energy of computation, that is the "
  "principle on which efficient hardware is built.")

h2("8.2 Future Scope")

bullets([
    ("Scaling the array", "extending the design from 4×4 to 16×16 or 32×32 "
     "would raise arithmetic intensity proportionally and reduce the relative "
     "cost of pipeline fill and drain."),
    ("Double-buffered weights", "adding a shadow weight register in each cell "
     "would allow the next weight matrix to be loaded while the current one is "
     "still computing, eliminating the reload overhead entirely."),
    ("Mixed and reduced precision", "supporting 4-bit weights or bfloat16 "
     "operands would allow accuracy to be traded against throughput on a "
     "per-layer basis."),
    ("Sparsity exploitation", "many trained networks contain a majority of "
     "zero weights; gating the multiplier when a weight is zero would reduce "
     "dynamic power substantially at very small area cost."),
    ("Deeper arithmetic pipelining", "inserting a register between the "
     "multiplier and the adder would shorten the critical path and raise the "
     "achievable clock frequency at the cost of one additional cycle of "
     "latency."),
    ("Integration with an activation pipeline", "adding ReLU, pooling and "
     "normalisation units at the array output would allow a complete network "
     "layer to be executed without host intervention."),
    ("FPGA and ASIC implementation", "synthesising the design to obtain real "
     "area, timing and power figures, and comparing them against the "
     "theoretical analysis presented here."),
])

# ===========================================================================
# REFERENCES
# ===========================================================================

h1_plain("References")

refs = [
    "H. T. Kung and C. E. Leiserson, \"Systolic Arrays (for VLSI),\" in "
    "Sparse Matrix Proceedings 1978, Society for Industrial and Applied "
    "Mathematics, 1979, pp. 256–282.",

    "H. T. Kung, \"Why Systolic Architectures?,\" IEEE Computer, vol. 15, "
    "no. 1, pp. 37–46, January 1982.",

    "N. P. Jouppi et al., \"In-Datacenter Performance Analysis of a Tensor "
    "Processing Unit,\" in Proceedings of the 44th Annual International "
    "Symposium on Computer Architecture (ISCA), Toronto, Canada, 2017, "
    "pp. 1–12.",

    "Y.-H. Chen, J. Emer and V. Sze, \"Eyeriss: A Spatial Architecture for "
    "Energy-Efficient Dataflow for Convolutional Neural Networks,\" in "
    "Proceedings of the 43rd Annual International Symposium on Computer "
    "Architecture (ISCA), Seoul, Korea, 2016, pp. 367–379.",

    "V. Sze, Y.-H. Chen, T.-J. Yang and J. S. Emer, \"Efficient Processing of "
    "Deep Neural Networks: A Tutorial and Survey,\" Proceedings of the IEEE, "
    "vol. 105, no. 12, pp. 2295–2329, December 2017.",

    "M. Horowitz, \"Computing's Energy Problem (and What We Can Do About It),\" "
    "in IEEE International Solid-State Circuits Conference (ISSCC) Digest of "
    "Technical Papers, San Francisco, USA, 2014, pp. 10–14.",

    "J. L. Hennessy and D. A. Patterson, Computer Architecture: A Quantitative "
    "Approach, 6th ed. Cambridge, MA: Morgan Kaufmann, 2019.",

    "J. L. Hennessy and D. A. Patterson, \"A New Golden Age for Computer "
    "Architecture,\" Communications of the ACM, vol. 62, no. 2, pp. 48–60, "
    "February 2019.",

    "I. Goodfellow, Y. Bengio and A. Courville, Deep Learning. Cambridge, MA: "
    "MIT Press, 2016.",

    "N. Weste and D. Harris, CMOS VLSI Design: A Circuits and Systems "
    "Perspective, 4th ed. Boston, MA: Addison-Wesley, 2010.",

    "IEEE Standard for SystemVerilog — Unified Hardware Design, Specification, "
    "and Verification Language, IEEE Std 1800-2017, 2018.",

    "G. H. Golub and C. F. Van Loan, Matrix Computations, 4th ed. Baltimore, "
    "MD: Johns Hopkins University Press, 2013.",
]

for i, ref in enumerate(refs, 1):
    par = doc.add_paragraph()
    par.paragraph_format.left_indent = Inches(0.45)
    par.paragraph_format.first_line_indent = Inches(-0.45)
    par.paragraph_format.space_after = Pt(8)
    par.paragraph_format.line_spacing = 1.2
    r = par.add_run(f"[{i}]\t")
    r.bold = True
    par.add_run(ref)

# ===========================================================================

doc.save(OUTPUT)
print(f"Written: {OUTPUT}")
print(f"Chapters: {CHAPTER[0]}")
print(f"Figures:  {sum(_fig_no.values())}")
print(f"Tables:   {sum(_tab_no.values())}")
print(f"Equations:{sum(_eq_no.values())}")
print("Open in Word and press Ctrl+A then F9 to build the Table of Contents.")
