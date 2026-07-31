#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
generate_pe_module_report.py

Builds a short, two-column, topic-wise MODULE REPORT for the Processing
Element of a weight-stationary systolic array.

Layout:  page 1  = title page (single column)
         pages 2+ = two-column textbook style body (~5 pages)

Run:  python generate_pe_module_report.py
Out:  PE_Module_Report.docx
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUTPUT = "PE_Module_Report.docx"

BODY_FONT = "Times New Roman"
MONO_FONT = "Consolas"
ACCENT = RGBColor(0x1F, 0x3B, 0x73)
GREY = RGBColor(0x44, 0x44, 0x44)

# ---------------------------------------------------------------------------
# XML helpers
# ---------------------------------------------------------------------------


def set_columns(section, num=2, space_twips=340):
    """Turn a section into an n-column layout."""
    sectPr = section._sectPr
    cols = sectPr.find(qn("w:cols"))
    if cols is None:
        cols = OxmlElement("w:cols")
        sectPr.append(cols)
    cols.set(qn("w:num"), str(num))
    cols.set(qn("w:space"), str(space_twips))
    cols.set(qn("w:equalWidth"), "1")


def shade(par, fill="F4F6F9"):
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill)
    par._p.get_or_add_pPr().append(shd)


def shade_cell(cell, fill="1F3B73"):
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill)
    cell._tc.get_or_add_tcPr().append(shd)


def box(par, colour="9AA5B1"):
    pbdr = OxmlElement("w:pBdr")
    for edge in ("top", "left", "bottom", "right"):
        e = OxmlElement(f"w:{edge}")
        e.set(qn("w:val"), "single")
        e.set(qn("w:sz"), "6")
        e.set(qn("w:space"), "4")
        e.set(qn("w:color"), colour)
        pbdr.append(e)
    par._p.get_or_add_pPr().append(pbdr)


# ---------------------------------------------------------------------------
# Document set-up
# ---------------------------------------------------------------------------

doc = Document()

sec0 = doc.sections[0]
sec0.page_width = Inches(8.27)          # A4
sec0.page_height = Inches(11.69)
for s_attr, val in (("left_margin", 0.9), ("right_margin", 0.9),
                    ("top_margin", 0.9), ("bottom_margin", 0.9)):
    setattr(sec0, s_attr, Inches(val))

normal = doc.styles["Normal"]
normal.font.name = BODY_FONT
normal.font.size = Pt(10.5)
normal.element.rPr.rFonts.set(qn("w:eastAsia"), BODY_FONT)
normal.paragraph_format.line_spacing = 1.18
normal.paragraph_format.space_after = Pt(6)
normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

h2s = doc.styles["Heading 2"]
h2s.font.name = BODY_FONT
h2s.font.size = Pt(11)
h2s.font.bold = True
h2s.font.color.rgb = ACCENT
h2s.paragraph_format.space_before = Pt(9)
h2s.paragraph_format.space_after = Pt(3)
h2s.paragraph_format.line_spacing = 1.0
h2s.paragraph_format.keep_with_next = True

h3s = doc.styles["Heading 3"]
h3s.font.name = BODY_FONT
h3s.font.size = Pt(10)
h3s.font.bold = True
h3s.font.italic = True
h3s.font.color.rgb = ACCENT
h3s.paragraph_format.space_before = Pt(6)
h3s.paragraph_format.space_after = Pt(2)
h3s.paragraph_format.line_spacing = 1.0
h3s.paragraph_format.keep_with_next = True

capstyle = doc.styles.add_style("Fig Caption", WD_STYLE_TYPE.PARAGRAPH)
capstyle.font.name = BODY_FONT
capstyle.font.size = Pt(8.5)
capstyle.font.italic = True
capstyle.font.color.rgb = GREY
capstyle.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
capstyle.paragraph_format.line_spacing = 1.0
capstyle.paragraph_format.space_after = Pt(8)

# ---------------------------------------------------------------------------
# Content helpers
# ---------------------------------------------------------------------------

_n = {"fig": 0, "tab": 0, "eq": 0}


def topic(text):
    return doc.add_heading(text, level=2)


def sub(text):
    return doc.add_heading(text, level=3)


def p(text):
    par = doc.add_paragraph()
    par.add_run(text)
    return par


def bullets(items):
    for it in items:
        par = doc.add_paragraph(style="List Bullet")
        par.paragraph_format.line_spacing = 1.05
        par.paragraph_format.space_after = Pt(2)
        par.paragraph_format.left_indent = Inches(0.20)
        if isinstance(it, tuple):
            par.add_run(it[0]).bold = True
            par.add_run(" — " + it[1])
        else:
            par.add_run(it)


def equation(text):
    _n["eq"] += 1
    par = doc.add_paragraph()
    par.alignment = WD_ALIGN_PARAGRAPH.CENTER
    par.paragraph_format.space_before = Pt(5)
    par.paragraph_format.space_after = Pt(5)
    par.paragraph_format.line_spacing = 1.0
    r = par.add_run(text)
    r.italic = True
    r.font.size = Pt(10.5)
    return par


def figure(art, caption, pt=7.0):
    _n["fig"] += 1
    par = doc.add_paragraph()
    par.alignment = WD_ALIGN_PARAGRAPH.CENTER
    par.paragraph_format.line_spacing = 0.92
    par.paragraph_format.space_before = Pt(7)
    par.paragraph_format.space_after = Pt(2)
    r = par.add_run(art)
    r.font.name = MONO_FONT
    r.font.size = Pt(pt)
    r._element.rPr.rFonts.set(qn("w:eastAsia"), MONO_FONT)
    box(par)
    shade(par, "FAFBFC")

    c = doc.add_paragraph(style="Fig Caption")
    c.add_run(f"Fig. {_n['fig']}  {caption}")
    return _n["fig"]


def table(caption, headers, rows, widths=None, size=8.0):
    _n["tab"] += 1
    c = doc.add_paragraph(style="Fig Caption")
    c.paragraph_format.space_before = Pt(7)
    c.paragraph_format.space_after = Pt(2)
    c.add_run(f"Table {_n['tab']}  {caption}")

    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.autofit = False

    hdr = t.rows[0].cells
    for i, htxt in enumerate(headers):
        shade_cell(hdr[i])
        par = hdr[i].paragraphs[0]
        par.alignment = WD_ALIGN_PARAGRAPH.CENTER
        par.paragraph_format.line_spacing = 1.0
        par.paragraph_format.space_after = Pt(1)
        r = par.add_run(htxt)
        r.bold = True
        r.font.size = Pt(size)
        r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            par = cells[i].paragraphs[0]
            par.alignment = (WD_ALIGN_PARAGRAPH.LEFT if i == 0
                             else WD_ALIGN_PARAGRAPH.CENTER)
            par.paragraph_format.line_spacing = 1.0
            par.paragraph_format.space_after = Pt(1)
            r = par.add_run(str(val))
            r.font.size = Pt(size)

    if widths:
        for row in t.rows:
            for i, w in enumerate(widths):
                row.cells[i].width = Inches(w)

    doc.add_paragraph().paragraph_format.space_after = Pt(4)
    return _n["tab"]


# ===========================================================================
# PAGE 1 — TITLE PAGE (single column)
# ===========================================================================

for _ in range(4):
    doc.add_paragraph()

par = doc.add_paragraph()
par.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = par.add_run("MODULE REPORT")
r.bold = True
r.font.size = Pt(13)
r.font.color.rgb = GREY

par = doc.add_paragraph()
par.alignment = WD_ALIGN_PARAGRAPH.CENTER
par.paragraph_format.space_before = Pt(20)
par.paragraph_format.line_spacing = 1.25
r = par.add_run("PROCESSING ELEMENT\nFOR A WEIGHT-STATIONARY\nSYSTOLIC ARRAY")
r.bold = True
r.font.size = Pt(24)
r.font.color.rgb = ACCENT

par = doc.add_paragraph()
par.alignment = WD_ALIGN_PARAGRAPH.CENTER
par.paragraph_format.space_before = Pt(14)
r = par.add_run("Design and Operation of a Single 8-bit\n"
                "Multiply–Accumulate Processing Element")
r.italic = True
r.font.size = Pt(12)

par = doc.add_paragraph()
par.alignment = WD_ALIGN_PARAGRAPH.CENTER
par.paragraph_format.space_before = Pt(16)
r = par.add_run("─" * 34)
r.font.color.rgb = ACCENT

for _ in range(3):
    doc.add_paragraph()

par = doc.add_paragraph()
par.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = par.add_run("Submitted by")
r.font.size = Pt(11)
r.font.color.rgb = GREY

for name in ("[ Name 1 ]", "[ Name 2 ]", "[ Name 3 ]"):
    par = doc.add_paragraph()
    par.alignment = WD_ALIGN_PARAGRAPH.CENTER
    par.paragraph_format.space_after = Pt(2)
    r = par.add_run(name)
    r.bold = True
    r.font.size = Pt(12.5)

doc.add_paragraph()

par = doc.add_paragraph()
par.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = par.add_run("Guide")
r.font.size = Pt(11)
r.font.color.rgb = GREY

par = doc.add_paragraph()
par.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = par.add_run("[ Guide Name ]")
r.bold = True
r.font.size = Pt(12)

for _ in range(4):
    doc.add_paragraph()

for txt, bold, size in (("[ Department ]", True, 11.5),
                        ("[ Institution ]", True, 11.5),
                        ("[ Month, Year ]", False, 11)):
    par = doc.add_paragraph()
    par.alignment = WD_ALIGN_PARAGRAPH.CENTER
    par.paragraph_format.space_after = Pt(2)
    r = par.add_run(txt)
    r.bold = bold
    r.font.size = Pt(size)

# ===========================================================================
# PAGES 2+ — TWO-COLUMN BODY
# ===========================================================================

sec1 = doc.add_section(WD_SECTION.NEW_PAGE)
sec1.page_width = Inches(8.27)
sec1.page_height = Inches(11.69)
sec1.left_margin = Inches(0.75)
sec1.right_margin = Inches(0.75)
sec1.top_margin = Inches(0.85)
sec1.bottom_margin = Inches(0.85)
set_columns(sec1, 2)

# --- Overview --------------------------------------------------------------

topic("1.  Overview")

p("A Processing Element (PE) is the basic computing cell of a systolic array. "
  "It is a small synchronous circuit that performs one multiply–accumulate "
  "(MAC) operation per clock cycle and then passes its data on to its "
  "neighbouring cells. A complete systolic array contains nothing but copies "
  "of this one cell wired together, so the behaviour of the whole machine "
  "follows entirely from the behaviour of a single PE.")

p("The PE described here is designed for a weight-stationary array. It holds "
  "one 8-bit signed weight in an internal register for the whole of a "
  "computation, multiplies that weight by an activation arriving from its "
  "left neighbour, adds the product to a partial sum arriving from the cell "
  "above, and passes the result to the cell below. It contains no instruction "
  "memory, no decoder and no addressable storage — its behaviour is fixed by "
  "its wiring and modulated only by two control signals.")

# --- Role ------------------------------------------------------------------

topic("2.  Role in the Systolic Array")

p("Matrix multiplication is defined element-wise as")

equation("Cᵢⱼ = Σ Aᵢₖ · Bₖⱼ")

p("which is a sum of products — exactly one MAC per term. A systolic array "
  "maps this directly onto hardware: each PE evaluates one term of the sum, "
  "and the summation itself is performed by physically passing the partial "
  "sum from one PE down to the next. No accumulator is ever addressed and no "
  "intermediate result is written to memory.")

p("Three streams therefore meet inside every PE. The weight is stationary and "
  "stays in the cell. The activation moves horizontally, entering from the "
  "left and leaving to the right. The partial sum moves vertically, entering "
  "from the top and leaving at the bottom. Because the streams travel in "
  "different directions, an operand meets a different partner in every cell "
  "it visits — which is how a single value fetched once from memory "
  "contributes to many different results.")

# --- MAC -------------------------------------------------------------------

topic("3.  The Multiply–Accumulate Operation")

p("The single arithmetic operation performed by the PE is")

equation("Psum = Psum + (A × B)")

p("or, in terms of the actual signal names of this module,")

equation("psum_out = psum_in + (weight_q × a_in)")

p("This is the complete functional specification of the cell. Everything "
  "else — the registers, the multiplexers, the control signals — exists only "
  "to deliver the right operands to this equation at the right time and to "
  "move the result to the right place. One PE performs one MAC per cycle; a "
  "4×4 array of them performs sixteen.")

# --- Interface -------------------------------------------------------------

topic("4.  Interface")

figure(
    "              psum_in [31:0]\n"
    "                    │\n"
    "                    ▼\n"
    "     ┌────────────────────────────┐\n"
    "     │                            │\n"
    " clk─►                            │\n"
    "rst_n─►     PROCESSING            ├─► valid_out\n"
    "wload─►        ELEMENT            │\n"
    "  v_in─►                          ├─► a_out [7:0]\n"
    "  a_in─►    IN_W = 8              │\n"
    " [7:0] │    ACC_W = 32            │\n"
    "     └─────────────┬──────────────┘\n"
    "                   │\n"
    "                   ▼\n"
    "             psum_out [31:0]",
    "Pin diagram of the processing element.")

table("Port description of the processing element.",
      ["Signal", "Dir / Width", "Function"],
      [["clk", "in, 1", "Global clock; all registers update on the rising edge"],
       ["rst_n", "in, 1", "Active-low synchronous reset; clears all registers"],
       ["wload", "in, 1", "1 = weight-load phase, 0 = compute phase"],
       ["valid_in", "in, 1", "1 = a_in carries a valid activation this cycle"],
       ["a_in", "in, 8", "Signed activation from the left neighbour"],
       ["a_out", "out, 8", "Registered activation to the right neighbour"],
       ["valid_out", "out, 1", "Registered valid flag to the right neighbour"],
       ["psum_in", "in, 32", "Signed partial sum from above; weight during wload"],
       ["psum_out", "out, 32", "Signed partial sum to below; weight during wload"]],
      widths=[0.62, 0.62, 1.71], size=7.5)

# --- Internal architecture -------------------------------------------------

topic("5.  Internal Architecture")

figure(
    "   psum_in[31:0]            a_in[7:0]\n"
    "        │                       │\n"
    "   ┌────┴─────┬─────────┐       ├──────┐\n"
    "   │ [7:0]    │         │       │      ▼\n"
    "   ▼          │         │       │  ┌───────┐\n"
    "┌────────┐    │         │       │  │  a_q  │\n"
    "│weight_q│◄── wload     │       │  └───┬───┘\n"
    "└───┬────┘              │       │      ▼\n"
    "    ├──────────┐        │       │    a_out\n"
    "    ▼          │        │       │\n"
    "┌─────────────────┐     │       │\n"
    "│  MUL  8×8 → 16  │◄────┼───────┘\n"
    "└────────┬────────┘     │\n"
    "         │ sign-extend  │\n"
    "         ▼              │\n"
    "┌─────────────────┐     │\n"
    "│   ADD  32-bit   │◄────┤\n"
    "└────────┬────────┘     │\n"
    "         ▼              │\n"
    "┌─────────────────┐     │\n"
    "│ MUX ◄ valid_in  │◄────┘\n"
    "└────────┬────────┘\n"
    "         ▼\n"
    "┌─────────────────┐\n"
    "│     psum_q      │\n"
    "└────────┬────────┘\n"
    "         ▼\n"
    "┌─────────────────┐\n"
    "│  MUX ◄ wload    │◄── weight_q (sign-extended)\n"
    "└────────┬────────┘\n"
    "         ▼\n"
    "   psum_out[31:0]\n"
    "\n"
    "  valid_in ──►[ valid_q ]──► valid_out",
    "Internal datapath of the processing element.", pt=6.8)

p("The cell contains four registers and three combinational blocks. There is "
  "exactly one register between any input and any output, so the PE adds "
  "exactly one cycle of latency to every path through it.")

sub("weight_q — the stationary weight")

p("An 8-bit register that loads from psum_in[7:0] on any rising clock edge "
  "on which wload is asserted, and holds its value otherwise. Because the "
  "register has an explicit enable, the compute traffic that later passes "
  "over the same psum wires cannot disturb it. This hold behaviour is what "
  "makes the architecture weight-stationary, and it is the single most "
  "important property of the cell. Only the low eight bits of the 32-bit bus "
  "are captured; the upper twenty-four are ignored.")

sub("a_q and valid_q — the activation pipeline")

p("a_q captures a_in on every clock edge and drives a_out; valid_q does the "
  "same for valid_in. Their purpose is to break the horizontal combinational "
  "path, so that an activation advances exactly one PE per cycle regardless "
  "of how wide the array is. valid_q travels alongside a_q so that the "
  "qualifier always arrives with the activation it qualifies.")

p("Note that the multiplier is driven by a_in, the incoming value, not by "
  "a_q. The PE multiplies the activation present on its input during the "
  "current cycle and simultaneously registers that same value for its "
  "neighbour. Multiplying a_q instead would delay the arithmetic by one cycle "
  "relative to the data movement and destroy the alignment between "
  "activations and partial sums.")

sub("psum_q — the accumulator")

p("A 32-bit register that loads psum_in + weight_q × a_in when valid_in is "
  "high, and psum_in unchanged when it is low. It never holds its previous "
  "value: the accumulation is not local to the cell but distributed down the "
  "column, with each PE adding its contribution to a sum that is passing "
  "through. This is what allows a new independent dot product to enter the "
  "column on every single cycle.")

# --- Weight loading --------------------------------------------------------

topic("6.  Weight Loading and the Shared psum Path")

p("The psum wires carry two entirely different kinds of information at two "
  "different times. During the weight-load phase (wload = 1) they carry "
  "weights being shifted down the column; during the compute phase "
  "(wload = 0) they carry partial sums. The output multiplexer selects "
  "accordingly: it drives the sign-extended contents of weight_q when wload "
  "is high, and psum_q when wload is low.")

p("The benefit is architectural rather than merely economical. Because "
  "weight_q drives psum_out during a load, each column of the array behaves "
  "as a vertical shift register, and all the weights of a column can be "
  "pushed in one per cycle through a single port at the top of the array. An "
  "n × n array therefore needs only n input ports to load n² weights. Without "
  "this time-sharing, weight distribution would dominate the wiring of the "
  "chip.")

p("The sign extension in this path matters. A weight of −5, held as 8'hFB, "
  "must appear on the 32-bit bus as 32'hFFFF_FFFB and not as 32'h0000_00FB; "
  "otherwise the PE below would receive +251 instead of −5.")

# --- Operation -------------------------------------------------------------

topic("7.  Cycle-by-Cycle Operation")

p("The behaviour of the cell is best seen through an example. The weight −3 "
  "is loaded, after which activations arrive with incoming partial sums.")

table("Operation of one PE with weight −3 loaded in cycle 0.",
      ["Cyc", "wl", "v", "a_in", "psum_in", "psum_out after edge"],
      [["0", "1", "0", "—", "−3", "−3  (weight passes down)"],
       ["1", "0", "1", "4", "100", "100 + (−3×4) = 88"],
       ["2", "0", "1", "5", "200", "200 + (−3×5) = 185"],
       ["3", "0", "0", "9", "300", "300  (unchanged)"],
       ["4", "0", "1", "6", "300", "300 + (−3×6) = 282"]],
      widths=[0.26, 0.24, 0.20, 0.34, 0.46, 1.45], size=7.5)

p("Cycle 0 loads the weight, and psum_out carries the weight itself so that "
  "the PE below can capture it on the next cycle. Cycles 1 and 2 are normal "
  "MAC operations. Cycle 3 shows the importance of the valid gate: although "
  "a_in carries the value 9, valid_in is low, so the partial sum passes "
  "through untouched. Cycle 4 resumes normal operation. Throughout, weight_q "
  "never changes — the weight is stationary.")

figure(
    "         │ 0 │ 1 │ 2 │ 3 │ 4 │\n"
    "  clk   ─┐ ┌─┐ ┌─┐ ┌─┐ ┌─┐ ┌─\n"
    "         └─┘ └─┘ └─┘ └─┘ └─┘\n"
    "\n"
    "  rst_n ──────────────────────\n"
    "\n"
    "  wload ──┐\n"
    "          └───────────────────\n"
    "\n"
    "  v_in  ────┐       ┌─┐\n"
    "            └───────┘ └───────\n"
    "\n"
    "  a_in  ──╳ ─ ╳ 4 ╳ 5 ╳ 9 ╳ 6 ╳\n"
    "\n"
    "  psum_in ╳ −3╳100╳200╳300╳300╳\n"
    "\n"
    "  psum_out╳ −3╳ 88╳185╳300╳282╳\n"
    "           ▲   ▲       ▲\n"
    "        weight MAC   bubble\n"
    "         out        (v_in=0)",
    "Waveform of the sequence in Table 2. Outputs change one cycle after the "
    "inputs that produce them.", pt=6.8)

# --- Neighbours ------------------------------------------------------------

topic("8.  Connection to Neighbouring Cells")

figure(
    "    psum_in        psum_in\n"
    "       │              │\n"
    "       ▼              ▼\n"
    "   ┌────────┐    ┌────────┐\n"
    "──►│ PE i,j ├───►│PE i,j+1├──►\n"
    "   └───┬────┘    └───┬────┘\n"
    "       ▼             ▼\n"
    "   ┌────────┐    ┌────────┐\n"
    "──►│PE i+1,j├───►│PE i+1,  ├──►\n"
    "   └───┬────┘    └───┬────┘\n"
    "       ▼             ▼\n"
    "   psum_out       psum_out",
    "Nearest-neighbour interconnection between cells.")

bullets([
    ("From the left", "a_in and valid_in come from the a_out and valid_out of "
     "the PE to the left; cells in column 0 take them from the array boundary."),
    ("To the right", "a_out and valid_out drive the next PE. In the last "
     "column they are unconnected — the activation has been fully used."),
    ("From above", "psum_in comes from the psum_out of the PE above; the top "
     "row takes it from the array boundary (zero, or a weight during load)."),
    ("To below", "psum_out drives the PE below. In the bottom row it is the "
     "completed dot product leaving the array."),
])

p("Every one of these connections is registered inside the cell, so the wire "
  "between two adjacent PEs is the only combinational path between them. This "
  "is what allows the array to be enlarged without any reduction in clock "
  "frequency.")

# --- Design considerations -------------------------------------------------

topic("9.  Design Considerations")

sub("Signed arithmetic")

p("Both operands are 8-bit two's-complement values in the range −128 to +127, "
  "so the multiplier must sign-extend them before forming partial products. "
  "Treating a negative weight as unsigned would interpret −1 as +255 and "
  "corrupt the result completely.")

sub("Product and accumulator width")

p("The product of two 8-bit signed values needs 16 bits: the extreme case "
  "(−128) × (−128) = +16 384 does not fit in 15. The product is sign-extended "
  "to 32 bits before the add. The 32-bit accumulator is deliberately "
  "generous — it leaves 16 guard bits, enough for 2¹⁶ worst-case "
  "accumulations, far beyond what a small array needs.")

sub("Valid gating")

p("The MAC is performed only when valid_in is asserted; otherwise psum_in "
  "passes through unchanged. This is essential during pipeline fill and "
  "drain, when some cells have no meaningful activation. Without the gate, "
  "the array would accumulate the products of whatever residual values "
  "happened to sit on the activation wires and silently corrupt every result.")

sub("Synchronous reset")

p("rst_n is active-low and synchronous: all four registers clear on a rising "
  "clock edge, not immediately. A synchronous reset avoids the "
  "recovery/removal timing problems of an asynchronous one and costs nothing "
  "here, since the array is never required to reset without a running clock.")

sub("Critical path")

p("The longest combinational path is one 8×8 multiplier followed by one "
  "32-bit adder. It does not lengthen as the array grows, because every "
  "inter-cell connection is registered. If a higher clock frequency were "
  "needed, a pipeline register could be inserted between the multiplier and "
  "the adder at the cost of one extra cycle of latency.")

# --- Verification ----------------------------------------------------------

topic("10.  Functional Verification")

p("Because the whole array is a replication of this one cell, a fault in the "
  "PE appears everywhere at once. The module is therefore verified on its own, "
  "against a golden reference model that recomputes psum + w × a "
  "independently, before being instantiated. Each test below is written to "
  "catch one specific class of design fault.")

table("Verification categories and the fault each is designed to catch.",
      ["Test", "Fault it catches"],
      [["Reset", "Asynchronous reset, registers that fail to clear, X values escaping reset"],
       ["Weight load", "Capturing more than 8 bits of psum_in; zero-extending instead of sign-extending the weight"],
       ["Stationarity", "A weight disturbed by compute traffic on the shared psum bus"],
       ["Hop timing", "A combinational (0-cycle) or over-registered (2-cycle) activation path; a_out and valid_out skewed apart"],
       ["MAC", "Wrong operand pairing, missing accumulate, unsigned multiply, product truncated before the add"],
       ["Bubble", "An accumulator that ignores valid_in and MACs on every cycle"],
       ["Signed", "Zero-extension of operands; mishandling of a negative incoming partial sum"],
       ["Boundary", "Truncation at 127×127, (−128)×(−128) and (−128)×127; a narrow accumulator dropping upper bits"],
       ["Transitions", "A weight that arrives one cycle late; residue left in psum_q by a load cycle"]],
      widths=[0.66, 2.29], size=7.5)

p("Two conventions keep the testbench itself trustworthy. Stimulus is driven "
  "on the falling clock edge and outputs are sampled after the rising edge, "
  "so driving and checking can never race. And no expected value is ever "
  "typed by hand — every one is produced by the golden function, so a "
  "mistake in the testbench cannot accidentally agree with the same mistake "
  "in the design.")

p("The bubble test deserves particular mention. It drives deliberately "
  "non-zero junk (8'h7F) on a_in while valid_in is low, with a non-zero "
  "weight loaded. Using a_in = 0 instead would pass even if the valid gate "
  "were missing entirely, because the product would be zero either way — a "
  "test that cannot fail is worse than no test at all.")

# --- Summary ---------------------------------------------------------------

topic("11.  Summary")

par = doc.add_paragraph()
par.paragraph_format.left_indent = Inches(0.08)
par.paragraph_format.right_indent = Inches(0.08)
par.paragraph_format.space_before = Pt(3)
shade(par, "F4F6F9")
r = par.add_run(
    "The processing element is a small synchronous cell containing one 8×8 "
    "signed multiplier, one 32-bit adder and four registers: weight_q for the "
    "stationary coefficient, a_q and valid_q for the horizontal activation "
    "pipeline, and psum_q for the vertical partial-sum pipeline. It computes "
    "psum_out = psum_in + weight_q × a_in whenever valid_in is asserted and "
    "passes the partial sum through unchanged otherwise. weight_q loads only "
    "when wload is asserted, which is what makes the design "
    "weight-stationary, and the psum wires are time-shared between weight "
    "distribution and partial-sum accumulation so that an n × n array can be "
    "loaded through only n ports. Because every interface is registered, the "
    "cell adds exactly one cycle of latency in each direction and can be "
    "replicated indefinitely without lengthening the critical path.")
r.font.size = Pt(9.5)

# ===========================================================================

doc.save(OUTPUT)
print(f"Written: {OUTPUT}")
print(f"Figures: {_n['fig']}  Tables: {_n['tab']}  Equations: {_n['eq']}")
